"""
Generate Phase I vs Phase II comparison Word document — Unified Telecom Analyzer.
Summarizes what has been done in Phase I (complete) and what is planned for
Phase II (novelty + real-world implementation), per the M.Tech review structure.
"""

import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from generate_study_doc import (
    new_doc, h1, h2, body, bullet, comparison_table, key_point,
    add_horizontal_rule, NAVY, CYAN, GREEN, GREY,
)


def cover_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIFIED TELECOM ANALYZER")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Phase I vs Phase II — Status & Comparison")
    r2.font.size = Pt(18)
    r2.font.color.rgb = CYAN

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("What we have done (Phase I) and what we are going to do (Phase II)")
    r3.font.size = Pt(13)
    r3.font.color.rgb = GREY
    r3.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    add_horizontal_rule(doc)

    for line in [
        "Candidate: Vikas",
        "Programme: M.Tech Data Science — PES Bangalore with Great Learning",
        "Project: Unified Telecom Analyzer (telecom-analyzer)",
        "Phase I status: COMPLETE",
        "Phase II status: PLANNING",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)
        r.font.color.rgb = GREY

    doc.add_page_break()


def ch_intro(doc):
    h1(doc, "1. Purpose of this Document")
    body(doc,
        "This document compares Phase I (Research Objective 1) — which is "
        "complete — with Phase II (Research Objective II), which focuses on "
        "novelty and real-world implementation. It records what has been "
        "built so far, maps it against the review checkpoints (Zeroth / "
        "First / Second / Third Review), and lays out the plan for Phase II.")
    key_point(doc, "Quick summary",
        "Phase I delivered a complete, working, tested pipeline (parsing -> "
        "detection -> correlation -> prediction -> explanation -> feedback). "
        "Phase II will add a novelty contribution and move the system toward "
        "real-world deployment, with a benchmark comparison against existing "
        "approaches.")


def ch_phase1_status(doc):
    h1(doc, "2. Phase I — Research Objective 1 (Status: COMPLETE)")
    body(doc,
        "Phase I required building a unified, multi-modal anomaly detection "
        "and explanation framework for 5G networks. All planned code has "
        "been implemented, integrated, tested, and demoed.")

    h2(doc, "2.1 Review-wise status")
    comparison_table(doc,
        ["Review", "Required deliverables", "Status in this project"],
        [
            ("Zeroth Review",
             "Title & Abstract, Introduction, Literature Survey, Proposed "
             "System, Timeline, References",
             "Done — covered in README architecture/overview and original proposal"),
            ("First Review",
             "Architectural design, Algorithms/Techniques, Expected "
             "outcomes, References, 30% code (Phase 1)",
             "Done & exceeded — full architecture, parsers + 6-detector "
             "PCAP/KPI ensembles built"),
            ("Second Review",
             "Detailed algo design, Contribution of candidate, Intermediate "
             "results, References, 80% code (Phase 1)",
             "Done & exceeded — 18-detector ensemble (PCAP+KPI+Stats), "
             "event correlation, Prophet/LSTM 4h prediction"),
            ("Third Review",
             "Overall design, Experimental results, Performance evaluation, "
             "Model comparison, References, 100% code + Demo, Journal paper draft",
             "Done — full pipeline, dashboard, REST API, 56 tests, reviewer "
             "PPTs. Journal paper draft remaining."),
        ],
        col_widths=[1.3, 2.7, 2.5])

    h2(doc, "2.2 What has been built (Phase I components)")
    items = [
        "Parsers — PCAP (NAS/NGAP/RRC/F1AP/E1AP/XnAP via pyshark + raw "
        "discriminator fallback), DU/CU stats (srsRAN/OAI/NIST/Parquet), "
        "KPI (Excel/CSV vs 3GPP KPI catalogue)",
        "Detection — 18-detector ensemble across PCAP, KPI, and Stats "
        "(Isolation Forest, One-Class SVM, LOF, Elliptic Envelope, "
        "Statistical/Cascade, LSTM-AE applied per domain)",
        "Correlation — cross-source event router linking PCAP + KPI + "
        "Stats anomalies by cell/time window",
        "Prediction — Prophet + LSTM forecasting anomalies 4 hours ahead, "
        "sharing the live-detection schema (state='predicted', lead_time_h)",
        "Explanation — FAISS + MiniLM RAG over 38 3GPP-spec chunks -> "
        "Ollama (phi3:mini) LLM explanations, with rule-based fallback",
        "MLOps loop — engineer feedback store (JSONL) -> nightly retrainer "
        "adjusting detector thresholds from observed false-positive rate",
        "Interfaces — Streamlit dashboard, FastAPI REST API, CLI orchestrator",
        "Validation — 56 automated tests (smoke + full pipeline), all passing",
    ]
    for it in items:
        bullet(doc, it)

    body(doc, "")
    key_point(doc, "Remaining for Phase I close-out",
        "Only the journal-paper draft remains — all code, integration, "
        "dashboard/API, and tests are complete and demoed.")


def ch_phase2_plan(doc):
    h1(doc, "3. Phase II — Research Objective II (Novelty + Real-World Implementation)")
    body(doc,
        "Phase II builds on the completed Phase I pipeline. Its focus, per "
        "the programme structure, is (a) a novelty contribution beyond "
        "Phase I, and (b) real-world implementation with a comparison "
        "against existing systems.")

    h2(doc, "3.1 Candidate novelty directions")
    rows = [
        ("Causal root-cause analysis",
         "Event router currently correlates anomalies by cell/time",
         "Add causal graphs (e.g., Granger causality / structural causal "
         "models) so the tool explains WHY one anomaly led to another, not "
         "just that they co-occurred"),
        ("Online / incremental learning",
         "Nightly batch retrainer (cron) adjusts detector params offline",
         "Move to streaming/online updates (e.g., incremental Isolation "
         "Forest) for near-real-time adaptation"),
        ("Federated / multi-gNB learning",
         "Single-pipeline, single-dataset training",
         "Train detectors across multiple gNBs without centralizing raw "
         "data — addresses real-world privacy/scale constraints"),
        ("Explainability (XAI) upgrade",
         "RAG retrieves 3GPP spec text; LLM produces qualitative explanation",
         "Add SHAP/LIME per-detector feature attribution feeding into the "
         "explanation, making root-cause reasoning quantitatively grounded"),
        ("Real-world deployment",
         "Synthetic generated PCAP/KPI/Stats test data only",
         "Containerize (Docker/K8s) and connect to a live gNB/CU/DU feed or "
         "public 5G testbed dataset, run near-real-time"),
        ("Benchmark / comparison study",
         "Internal method-comparison matrix only (own 18 detectors vs each other)",
         "Quantitative comparison of the 18-detector ensemble vs published "
         "anomaly-detection approaches on a common dataset — required for "
         "Third Review 'Comparison with Existing system'"),
    ]
    comparison_table(doc,
        ["Direction", "Phase I baseline", "Phase II proposal"],
        rows, col_widths=[1.6, 2.2, 2.7])

    key_point(doc, "Recommendation",
        "Prioritize 'Real-world deployment' + 'Benchmark/comparison study' "
        "as the core Phase II novelty — the Phase I detector/RAG/MLOps stack "
        "is already strong, so Phase II should focus on proving it works on "
        "real data and how it compares to existing methods. The other "
        "directions (causal RCA, online learning, federated learning, XAI) "
        "can be secondary enhancements if time permits.")

    h2(doc, "3.2 Phase II review plan (template)")
    comparison_table(doc,
        ["Review", "Required deliverables", "Plan"],
        [
            ("First Review",
             "Phase 1 review, Novelty proposal, Title & Abstract, Proposed "
             "System (Phase II), Algorithms/Techniques, Expected outcomes, "
             "References, 40% code (Phase II)",
             "Finalize novelty direction(s) with guide; draft proposed "
             "system additions (deployment architecture / benchmark "
             "datasets); begin implementation"),
            ("Second Review",
             "Modified algorithm design, Contribution of candidate, "
             "Intermediate results, References, 80% code (Phase II)",
             "Complete novelty implementation; run on real-world / "
             "benchmark data; collect intermediate results"),
            ("Third Review",
             "Overall design (Phase I + II), Experimental results, "
             "Performance evaluation, Comparison with existing system, "
             "References, 100% code + Demo, Journal paper publication proof",
             "Final demo on real-world setup; publish comparison results; "
             "submit journal paper covering Phase I + II"),
        ],
        col_widths=[1.3, 2.8, 2.4])


def ch_comparison_summary(doc):
    h1(doc, "4. Side-by-Side: What We Have Done vs What We Are Going To Do")
    comparison_table(doc,
        ["Area", "Phase I (Done)", "Phase II (Planned)"],
        [
            ("Data sources",
             "Synthetic PCAP, KPI, DU/CU stats (generator scripts)",
             "Real-world / live gNB feed or public 5G testbed dataset"),
            ("Detection",
             "18-detector ensemble (6 methods x PCAP/KPI/Stats)",
             "Same ensemble + benchmark comparison vs published methods"),
            ("Correlation",
             "Rule-based cross-source correlation (cell/time window)",
             "Optional: causal-graph based root-cause linking"),
            ("Retraining",
             "Nightly batch retrainer (cron) from feedback log",
             "Optional: online/incremental retraining"),
            ("Explanation",
             "RAG (FAISS+MiniLM) -> local LLM (Ollama phi3:mini)",
             "Optional: XAI feature-attribution grounding for explanations"),
            ("Deployment",
             "Local dev (Streamlit + FastAPI on WSL2, single machine)",
             "Containerized (Docker/K8s) real-world deployment"),
            ("Evaluation",
             "Internal 56-test suite + method-comparison matrix",
             "Quantitative benchmark vs existing systems on common dataset"),
            ("Documentation",
             "README, reviewer PPTs, study guide, 56 tests",
             "Journal paper (Phase I + II), updated reviewer PPTs"),
        ],
        col_widths=[1.3, 2.6, 2.6])


def ch_next_steps(doc):
    h1(doc, "5. Immediate Next Steps")
    bullet(doc, "Finalize Phase II novelty direction(s) with guide (shortlist in Section 3.1)")
    bullet(doc, "Draft Phase I journal paper (code/demo already complete)")
    bullet(doc, "Identify a real-world or public benchmark dataset for Phase II")
    bullet(doc, "Once direction is locked, fill in the Phase II review plan (Section 3.2) "
                "with concrete milestones and timeline")


def build():
    print("Building Phase I vs Phase II comparison document...")
    doc = new_doc()

    cover_page(doc)
    ch_intro(doc)
    ch_phase1_status(doc)
    doc.add_page_break()
    ch_phase2_plan(doc)
    doc.add_page_break()
    ch_comparison_summary(doc)
    ch_next_steps(doc)

    out_path = "/mnt/e/telecom-analyzer/docs/Phase1_vs_Phase2_Comparison.docx"
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Size: {os.path.getsize(out_path) // 1024} KB")


if __name__ == "__main__":
    build()
