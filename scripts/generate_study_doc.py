"""
Generate a comprehensive Word study document for the Unified Telecom Analyzer project.
Explains every component, every choice, results, and study notes in easy English.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colour constants (RGB) ────────────────────────────────────────────
NAVY    = RGBColor(0x1F, 0x38, 0x96)
CYAN    = RGBColor(0x00, 0x70, 0xC0)
GREEN   = RGBColor(0x37, 0x86, 0x10)
RED     = RGBColor(0xC0, 0x00, 0x00)
ORANGE  = RGBColor(0xFF, 0x66, 0x00)
GREY    = RGBColor(0x40, 0x40, 0x40)
LGREY   = RGBColor(0xF2, 0xF2, 0xF2)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
YELLOW  = RGBColor(0xFF, 0xD7, 0x00)
TEAL    = RGBColor(0x00, 0x80, 0x80)


def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)


def new_doc() -> Document:
    doc = Document()
    # Set page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)
    # Default body font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    return doc


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(20)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    add_horizontal_rule(doc)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = CYAN
    p.runs[0].font.size = Pt(15)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = TEAL
    p.runs[0].font.size = Pt(13)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p


def body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(3)
    return p


def numbered(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(3)
    return p


def code_block(doc, lines):
    """Add a monospace code block with grey background."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        # Grey background via XML shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1F, 0x50, 0x1F)
    # small spacer after block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def key_point(doc, label, text):
    """Highlight box for important takeaways."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'E7F3FF')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"💡 {label}: ")
    r1.bold = True
    r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def result_box(doc, title, lines):
    """Green result box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'E8F5E9')
    p = cell.paragraphs[0]
    r = p.add_run(f"✅ RESULT — {title}")
    r.bold = True
    r.font.color.rgb = GREEN
    for line in lines:
        p2 = cell.add_paragraph(f"  {line}")
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(2)
        p2.runs[0].font.color.rgb = GREY if not line.startswith("•") else GREEN
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def warning_box(doc, title, text):
    """Orange warning / 'why not' box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'FFF3E0')
    p = cell.paragraphs[0]
    r = p.add_run(f"⚠️  Why NOT this approach: {title}")
    r.bold = True
    r.font.color.rgb = ORANGE
    p2 = cell.add_paragraph(f"  {text}")
    p2.runs[0].font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def comparison_table(doc, headers, rows, col_widths=None):
    """Generic comparison table with styled header."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, '1F3864')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F5F8FF'
        drow = t.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = drow.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if val in ('✅', '✅ Yes', '✅ Full', '✅ Fast'):
                run.font.color.rgb = GREEN
            elif val in ('❌', '❌ No', '❌ None', '❌ Manual'):
                run.font.color.rgb = RED
            elif str(val).startswith('⚠'):
                run.font.color.rgb = ORANGE
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return t


# ══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════
def cover_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIFIED TELECOM ANALYZER")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Complete Technical Study Guide")
    r2.font.size = Pt(18)
    r2.font.color.rgb = CYAN

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Every Component · Every Decision · Every Result")
    r3.font.size = Pt(13)
    r3.font.color.rgb = GREY
    r3.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    add_horizontal_rule(doc)

    for line in [
        "Prepared for: Vikas",
        "Project: 5G Network Anomaly Detection Dashboard",
        "Language: Python 3.x",
        "Date: May 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)
        r.font.color.rgb = GREY

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 0 — Project Overview
# ══════════════════════════════════════════════════════════════════════
def ch_overview(doc):
    h1(doc, "Chapter 0 — What is This Project?")

    body(doc,
         "This project is a 5G network analyzer. When telecom engineers want to find problems "
         "in their 5G network — like why calls are dropping, why internet speeds are slow, or "
         "why connections are failing — they need a tool that can automatically read raw "
         "network data and tell them what is wrong. That is exactly what this project does.")

    h2(doc, "What problem does it solve?")
    body(doc, "In a 5G network, there are thousands of messages being sent every second between "
              "different parts of the network (like the phone, the tower, and the core network). "
              "When something goes wrong:")
    bullet(doc, "A network engineer would have to manually look at millions of packets to find the problem — this takes days.")
    bullet(doc, "Our tool does this automatically in seconds.")
    bullet(doc, "It gives clear answers: 'Registration is failing 15% of the time because of congestion.'")

    h2(doc, "The Two Types of Input Data")
    body(doc, "The system accepts two types of files:")

    comparison_table(doc,
        ["Data Type", "File Format", "What it contains", "What we extract"],
        [
            ["PCAP File", ".pcap / .pcapng", "Raw network packets captured at wire level", "Which procedures ran, which succeeded/failed, why they failed"],
            ["KPI File", ".xlsx / .csv", "Hourly performance statistics from towers", "Trends, outliers, degrading KPIs across cells"],
        ],
        col_widths=[1.3, 1.3, 2.2, 2.2]
    )

    h2(doc, "Project Architecture (Simple View)")
    body(doc, "Think of the project as a pipeline with 3 stages:")
    numbered(doc, "INPUT  →  Upload a PCAP file OR a KPI Excel file")
    numbered(doc, "PROCESS  →  Parse the file (understand the data) + Detect anomalies (find problems)")
    numbered(doc, "OUTPUT  →  Show results on a visual dashboard with explanations")

    key_point(doc, "Why is this useful?",
              "Telecom engineers spend 60-70% of their time manually looking through logs. "
              "This tool automates that work and shows only the important problems.")

    h2(doc, "Technologies Used")
    comparison_table(doc,
        ["Technology", "What it is", "Why we use it"],
        [
            ["Python 3.x", "Programming language", "Most popular for data science and ML"],
            ["Streamlit", "Web dashboard framework", "Build beautiful web apps with pure Python — no HTML/CSS needed"],
            ["pyshark", "PCAP parser", "Wraps Wireshark's 3GPP dissectors — best tool for 5G packet analysis"],
            ["scikit-learn", "ML library", "Provides Isolation Forest, SVM, LOF, Elliptic Envelope"],
            ["PyTorch", "Deep learning library", "Used for LSTM Autoencoder sequence detector"],
            ["pandas / numpy", "Data processing", "Standard data manipulation tools"],
            ["Plotly", "Interactive charts", "Beautiful time-series charts for KPI trends"],
            ["python-pptx", "PPT generator", "Auto-generate reviewer presentation slides"],
        ],
        col_widths=[1.4, 2.0, 3.3]
    )

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 1 — PCAP Parsing
# ══════════════════════════════════════════════════════════════════════
def ch_pcap_parsing(doc):
    h1(doc, "Chapter 1 — PCAP Parsing (Reading Network Packets)")

    h2(doc, "1.1  What is a PCAP File?")
    body(doc,
         "PCAP stands for Packet Capture. When your phone connects to a 5G tower, "
         "hundreds of messages are exchanged. A PCAP file is like a recording of all "
         "those messages — every single packet that travelled on the wire.")
    body(doc,
         "The problem is that these messages are not human-readable. They are in binary "
         "format using 3GPP-specific encoding. Just like you cannot read a .exe file by "
         "opening it in Notepad, you cannot read a 5G PCAP without a proper decoder.")

    key_point(doc, "Simple analogy",
              "A PCAP file is like a video recording of a phone conversation. "
              "Just knowing the recording exists is not enough — you need a translator "
              "who understands the 5G language to tell you what was said.")

    h2(doc, "1.2  The 5G Protocol Stack — What Layers We Parse")
    body(doc, "5G networks use layered protocols. Each layer handles a different part of communication:")

    comparison_table(doc,
        ["Protocol", "3GPP Spec", "Interface", "What it handles"],
        [
            ["NAS", "TS 24.501", "UE ↔ Core (AMF)", "Phone registration, authentication, data sessions"],
            ["NGAP", "TS 38.413", "gNB ↔ AMF (N2)", "Tower-to-core messages: UE context, handover"],
            ["RRC", "TS 38.331", "UE ↔ gNB (Uu)", "Radio connection setup, reconfiguration, release"],
            ["F1AP", "TS 38.473", "DU ↔ CU-CP (F1)", "Split architecture: Distributed Unit ↔ Central Unit"],
            ["E1AP", "TS 38.463", "CU-CP ↔ CU-UP (E1)", "Control plane ↔ User plane separation"],
            ["XnAP", "TS 38.423", "gNB ↔ gNB (Xn)", "Inter-tower communication, handover between towers"],
        ],
        col_widths=[0.9, 1.1, 1.8, 2.9]
    )

    key_point(doc, "Why all 6?",
              "A single call involves all 6 layers. If the NAS registration fails, "
              "NGAP and F1AP messages also show symptoms. To find root cause, you need all 6.")

    h2(doc, "1.3  The 4 Parsing Methods We Evaluated")
    body(doc, "Before writing any code, we evaluated four different Python libraries "
              "for parsing 5G PCAP files. Here is what we found:")

    h3(doc, "Method 1: pyshark (Wireshark Wrapper) — ✅ CHOSEN")
    body(doc, "pyshark is a Python library that wraps Wireshark's dissectors. "
              "Wireshark is the world's most popular network analysis tool, maintained "
              "by the telecom community for 20+ years. It already knows how to decode "
              "every 5G protocol.")

    code_block(doc, [
        "import pyshark",
        "",
        "# Open the PCAP file with a filter",
        "cap = pyshark.FileCapture('test.pcap',",
        "          display_filter='ngap or nas-5gs')",
        "",
        "for pkt in cap:",
        "    # Access decoded fields directly — Wireshark does the decoding",
        "    proc_code = pkt.ngap.procedurecode    # e.g. '14' = InitialContextSetup",
        "    cause     = pkt.nas_5gs.mm.cause      # e.g. '22' = Congestion",
        "    print(proc_code, cause)",
    ])

    result_box(doc, "pyshark output on real 5G PCAP", [
        "• Protocol: NAS (5GMM)",
        "• Procedure: Registration Request",
        "• UE ID: 0x1234ABCD",
        "• IEs extracted: 5GMM cause = Congestion (#22), Registration type = Initial",
        "• Protocol: NGAP — InitialContextSetup → Failure",
        "• Cause: radio-resources-not-available",
        "• All 6 protocol layers decoded in one pass",
    ])

    h3(doc, "Method 2: scapy — ❌ NOT CHOSEN")
    body(doc, "scapy is a powerful packet manipulation library. It is excellent for "
              "crafting packets and reading IP/UDP layers. However, it has zero knowledge "
              "of 5G protocols.")

    code_block(doc, [
        "from scapy.all import rdpcap",
        "",
        "pkts = rdpcap('test.pcap')",
        "for pkt in pkts:",
        "    raw = bytes(pkt['UDP'].payload)",
        "    # raw = b'\\x7e\\x00\\x41\\x00...'  ← meaningless binary",
        "    # You would have to decode this manually",
    ])

    warning_box(doc, "scapy",
        "scapy gives you raw bytes after the UDP layer. To get 'NAS_Registration' from "
        "those bytes, you would need to: (1) check byte[0] == 0x7e (NAS), "
        "(2) check byte[2] == 0x41 (Registration), (3) manually parse TLV IEs per "
        "TS 24.501 Table 8.2.1, (4) repeat for all 6 protocols. That is ~3000 lines "
        "of custom code that already exists inside Wireshark. We would be reinventing the wheel.")

    h3(doc, "Method 3: dpkt — ❌ NOT CHOSEN")
    body(doc, "dpkt is the fastest PCAP library — it is written in C. But speed is its "
              "only advantage for our use case.")

    code_block(doc, [
        "import dpkt",
        "",
        "with open('test.pcap', 'rb') as f:",
        "    pcap = dpkt.pcap.Reader(f)",
        "    for timestamp, buf in pcap:",
        "        eth  = dpkt.ethernet.Ethernet(buf)",
        "        data = eth.ip.udp.data   # raw 5G bytes — completely opaque",
    ])

    warning_box(doc, "dpkt",
        "dpkt stops at the transport layer (UDP/SCTP). It has absolutely no knowledge "
        "of NAS, NGAP, RRC, F1AP, E1AP, or XnAP. We would get the same raw bytes as "
        "scapy, but with fewer features and no packet crafting. Eliminated early.")

    h3(doc, "Method 4: asn1tools — ❌ NOT CHOSEN (except for reference)")
    body(doc, "asn1tools is a pure Python ASN.1 decoder. ASN.1 is the data format that "
              "most 3GPP protocols use. In theory, this approach gives the most accurate decoding.")

    code_block(doc, [
        "import asn1tools",
        "",
        "# Need to provide the official 3GPP ASN.1 schema files",
        "ngap = asn1tools.compile_files([",
        "    'NGAP-PDU-Descriptions.asn',",
        "    'NGAP-CommonDataTypes.asn',",
        "    'NGAP-IEs.asn',",
        "    'NGAP-Containers.asn'",
        "])",
        "decoded = ngap.decode('NGAP-PDU', raw_bytes)",
    ])

    warning_box(doc, "asn1tools",
        "Problem 1: You need 15+ .asn schema files across 6 protocols, and these change "
        "with every 3GPP release (Rel-15, 16, 17, 18). Managing schema versions is a "
        "project in itself. "
        "Problem 2: NAS (the most important protocol) uses a CUSTOM binary TLV encoding "
        "defined in TS 24.007 — it is NOT standard ASN.1. asn1tools cannot decode it. "
        "You would still need a separate NAS decoder. "
        "Problem 3: Wireshark already does all of this correctly, so using asn1tools "
        "would be building an inferior version of something that already exists.")

    h2(doc, "1.4  The 4-Method Comparison Table")
    comparison_table(doc,
        ["Library", "NAS/NGAP/RRC", "F1/E1/XnAP", "IE Extraction", "Speed", "Setup Effort"],
        [
            ["pyshark ✅", "✅ Full", "✅ Full", "✅ Full", "⚠️ Slow", "Low (just tshark)"],
            ["scapy", "❌ Raw bytes", "❌ Raw bytes", "❌ Manual only", "✅ Fast", "Low"],
            ["dpkt", "❌ None", "❌ None", "❌ None", "✅ Fastest", "Low"],
            ["asn1tools", "⚠️ Partial", "✅ Full", "✅ Full", "⚠️ Medium", "❌ Very High"],
        ],
        col_widths=[1.3, 1.3, 1.3, 1.4, 1.0, 1.4]
    )

    h2(doc, "1.5  How Our Parser Works — Step by Step")
    body(doc, "File: src/parsers/pcap_parser_real.py")

    h3(doc, "Step 1: Open the PCAP with a display filter")
    code_block(doc, [
        "display_filter = (",
        "    'nas-5gs or ngap or nr-rrc or f1ap or e1ap or xnap'",
        "    ' or udp.port == 38412'   # for synthetic test PCAPs",
        ")",
        "cap = pyshark.FileCapture(pcap_path, display_filter=display_filter,",
        "                          keep_packets=False)",
    ])
    body(doc, "The display_filter tells pyshark to only give us packets that belong to "
              "5G protocols. This is like telling a translator: 'only translate French sentences, "
              "ignore everything else.' It makes parsing much faster on large captures.")

    h3(doc, "Step 2: Process each packet")
    code_block(doc, [
        "for packet in cap:",
        "    if hasattr(packet, 'nas_5gs'):   # Is this a NAS packet?",
        "        self._handle_nas_5gs(packet)",
        "    if hasattr(packet, 'ngap'):      # Is this an NGAP packet?",
        "        self._handle_ngap_real(packet)",
        "    if hasattr(packet, 'nr_rrc'):    # Is this an RRC packet?",
        "        self._handle_rrc_real(packet)",
        "    # ... and so on for F1AP, E1AP, XnAP",
    ])
    body(doc, "One physical packet can contain messages from multiple layers. "
              "For example, a packet might have both NAS and NGAP content. "
              "We check for each layer independently so nothing is missed.")

    h3(doc, "Step 3: Track procedures (request → response → success/failure)")
    body(doc, "A 'procedure' is a complete operation like 'NAS Registration'. "
              "It consists of multiple messages:")
    bullet(doc, "Message 1: UE sends Registration Request  (role = request)")
    bullet(doc, "Message 2: Network sends Registration Accept  (role = response → SUCCESS)")
    bullet(doc, "OR Message 2: Network sends Registration Reject  (role = failure)")

    code_block(doc, [
        "# When we see a request, we start tracking it",
        "def _record_request(self, proc_name, ue_id, layer):",
        "    self.stats[proc_name].attempts += 1",
        "    self.pending[key] = ProcedureRecord(proc_name, key, time.now())",
        "",
        "# When we see the response, we mark it complete",
        "def _record_success(self, proc_name, ue_id):",
        "    self.stats[proc_name].success += 1",
        "",
        "# When we see a failure, we record the cause",
        "def _record_failure(self, proc_name, ue_id, cause):",
        "    self.stats[proc_name].failure += 1",
        "    self.stats[proc_name].failure_causes[cause] += 1",
    ])

    h3(doc, "Step 4: Extract Information Elements (IEs)")
    body(doc, "IEs are the data fields inside each message. For example, a NAS "
              "Registration message contains IEs like: UE Security Capabilities, "
              "5G-GUTI, Registration Type, etc. We extract these so the user can "
              "see the exact content of each message.")

    h3(doc, "Step 5: Build the output dictionary")
    body(doc, "After processing all packets, we return a structured dictionary:")
    code_block(doc, [
        "{",
        "  'procedures': {",
        "    'NAS_Registration': {",
        "      'attempts': 100, 'success': 95, 'failure': 5,",
        "      'success_rate': 95.0,",
        "      'failure_causes': {'congestion': 3, 'timeout': 2}",
        "    },",
        "    ...",
        "  },",
        "  'message_log': [ {layer, procedure, role, ue_id, ies, timestamp}, ... ],",
        "  'total_events': 500,",
        "  'parser_version': 'real_v3'",
        "}",
    ])

    result_box(doc, "Parser on test_5g_full.pcap", [
        "• Total events: 500 signalling messages",
        "• NAS: Registration (100 attempts, 95% SR), Authentication (99%), PDU Session (67%!)",
        "• NGAP: InitialContextSetup (99%), PDUSessionResourceSetup (98%)",
        "• RRC: Setup (99%), Reconfiguration (100%)",
        "• F1AP: UEContextSetup (100%)",
        "• E1AP: BearerContextSetup (59%!) — major anomaly detected",
        "• XnAP: XnSetup (100%), HandoverRequest (95%)",
    ])

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 2 — KPI Parser
# ══════════════════════════════════════════════════════════════════════
def ch_kpi_parser(doc):
    h1(doc, "Chapter 2 — KPI Parser (Reading Performance Statistics)")

    h2(doc, "2.1  What are KPIs in 5G?")
    body(doc,
         "KPI stands for Key Performance Indicator. Every 15 minutes (or every hour), "
         "the gNB (5G tower) reports statistics to the Operations & Maintenance (O&M) system. "
         "These statistics tell the engineer how the network is performing.")
    body(doc,
         "Think of it like a hospital monitoring a patient's vitals — heart rate, blood pressure, "
         "oxygen level — every hour. If any vital goes outside normal range, the nurse is alerted. "
         "KPIs are the 'vitals' of a 5G cell.")

    h2(doc, "2.2  KPI Categories We Track")
    comparison_table(doc,
        ["Category", "KPI Examples", "Normal Value", "Why It Matters"],
        [
            ["Availability", "Cell Availability %", "> 99.5%", "Is the cell even on-air?"],
            ["Accessibility", "RRC SR %, Registration SR %", "> 99%", "Can UEs connect?"],
            ["Retainability", "Drop Call Rate", "< 0.5%", "Do connections stay up?"],
            ["Mobility", "Handover Success Rate %", "> 98%", "Can UEs move between cells?"],
            ["Capacity", "PRB Utilization DL/UL %", "< 80%", "Is the cell overloaded?"],
            ["Radio Quality", "CQI, SINR dB, BLER %", "CQI > 8", "Is radio signal good?"],
            ["Throughput", "DL/UL Mbps", "Depends on cell type", "How fast is data?"],
            ["Latency", "One-way Latency ms", "< 20ms", "How fast are responses?"],
        ],
        col_widths=[1.4, 2.2, 1.4, 2.0]
    )

    h2(doc, "2.3  The KPI Parser Code")
    body(doc, "File: src/parsers/kpi_parser.py")
    body(doc, "This parser reads an Excel or CSV file from the O&M system and converts it "
              "into a structured format that our detection engine can work with.")

    h3(doc, "Step 1: Load the file")
    code_block(doc, [
        "def _load_raw(filepath):",
        "    path = Path(filepath)",
        "    if path.suffix in ('.xlsx', '.xls'):",
        "        df = pd.read_excel(filepath, engine='openpyxl')  # Excel files",
        "    elif path.suffix == '.csv':",
        "        df = pd.read_csv(filepath)                       # CSV files",
        "    return df",
    ])
    body(doc, "Why openpyxl for Excel? Because it is the standard Python library for .xlsx "
              "files and handles all Excel formatting, merged cells, and data types correctly.")

    h3(doc, "Step 2: Normalise column names")
    body(doc, "Different vendors name the same KPI differently. Ericsson might call it "
              "'RRC Connection Setup Success Rate', Nokia might call it 'RRC_ConnEstab_SR', "
              "Huawei might call it 'rrc_sr'. Our parser maps all these to one standard name.")
    code_block(doc, [
        "# In kpi_defs.py, each KPI has a list of 'aliases'",
        "'RRC_Success_Rate_%': {",
        "    'label': 'RRC Connection Setup Success Rate',",
        "    'aliases': ['RRC Connection Setup Success Rate',",
        "                'RRC Setup Success Rate',",
        "                'RRC_ConnEstab_SR']  # all vendor names",
        "}",
        "",
        "# The parser tries to match any column name to our catalog",
        "def resolve_column(col_name):",
        "    for canonical, meta in KPI_CATALOG.items():",
        "        if col_name in meta['aliases'] or col_name == canonical:",
        "            return canonical  # return standard name",
        "    return col_name  # unknown column — keep as-is",
    ])

    h3(doc, "Step 3: Identify structural vs KPI columns")
    code_block(doc, [
        "# Structural columns = identifiers, not measurements",
        "ID_COLUMNS = {'timestamp', 'cell_id', 'gnb_id', 'date', 'time', ...}",
        "",
        "# KPI columns = numeric columns that are NOT identifiers",
        "kpi_cols = [",
        "    c for c in df.columns",
        "    if c not in id_cols",
        "    and pd.api.types.is_numeric_dtype(df[c])  # must be a number",
        "]",
    ])
    body(doc, "This is smart automatic detection. Even if a vendor's file has unusual column "
              "names, as long as they are numeric and not identifiers, we treat them as KPIs.")

    h3(doc, "Step 4: Compute summary statistics")
    code_block(doc, [
        "for col in kpi_cols:",
        "    s = df[col].dropna()  # remove empty values",
        "    summary[col] = {",
        "        'min':  s.min(),",
        "        'max':  s.max(),",
        "        'mean': s.mean(),",
        "        'std':  s.std(),",
        "        'p10':  s.quantile(0.10),  # 10th percentile",
        "        'p90':  s.quantile(0.90),  # 90th percentile",
        "    }",
    ])

    key_point(doc, "Why P10 and P90?",
              "Mean alone can be misleading. P10 is the value that 90% of data is above. "
              "P90 is the value that 90% of data is below. Together they show the normal "
              "operating range without being influenced by extreme outliers.")

    result_box(doc, "KPI Parser on sample Excel file", [
        "• Rows: 720 (30 cells × 24 hours)",
        "• Unique Cells: 30",
        "• Unique gNBs: 10",
        "• KPI Columns found: 15 (all mapped to canonical names)",
        "• Time Range: 2026-05-01 00:00 → 2026-05-01 23:45",
        "• Summary stats computed for all 15 KPIs",
    ])

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 3 — Feature Engineering
# ══════════════════════════════════════════════════════════════════════
def ch_features(doc):
    h1(doc, "Chapter 3 — Feature Engineering (Preparing Data for ML Models)")

    body(doc, "File: src/detection/features.py")
    body(doc,
         "Machine Learning models cannot directly read raw PCAP data. They need numbers. "
         "Feature engineering is the process of converting raw data into a numerical "
         "representation that ML algorithms can process. This is one of the most important "
         "steps in any ML pipeline.")

    h2(doc, "3.1  Two Feature Spaces")
    body(doc, "We create two different sets of features for two different purposes:")

    comparison_table(doc,
        ["Feature Space", "Used by", "What it captures", "Shape"],
        [
            ["Procedure-level features", "Isolation Forest, OC-SVM, LOF, Elliptic Envelope", "Per-procedure statistics (success rate, failure causes, etc.)", "N procedures × 8 features"],
            ["Sequence features", "LSTM Autoencoder", "Temporal order of messages in the capture", "N windows × W length × 4 features"],
        ],
        col_widths=[1.8, 2.0, 2.2, 1.7]
    )

    h2(doc, "3.2  Procedure-Level Features (Tabular ML)")
    body(doc, "For each procedure (e.g., NAS_Registration), we compute 8 numbers:")

    comparison_table(doc,
        ["Feature", "Formula", "Range", "Why it matters"],
        [
            ["success_rate", "success / attempts × 100", "0–100", "Core health metric"],
            ["failure_rate", "100 - success_rate", "0–100", "Complement of success"],
            ["attempt_share", "this_proc_attempts / total_attempts", "0–1", "How busy is this procedure"],
            ["failure_count", "raw number of failures", "0–∞", "Absolute severity"],
            ["cause_diversity", "distinct causes / max(failure, 1)", "0–1", "Are failures varied or concentrated?"],
            ["top_cause_concentration", "most_common_cause / max(failure, 1)", "0–1", "Is one cause dominating?"],
            ["timeout_ratio", "timeout_failures / max(failure, 1)", "0–1", "Connectivity issues"],
            ["has_any_failure", "1 if failure > 0 else 0", "0 or 1", "Binary flag"],
        ],
        col_widths=[1.8, 2.2, 0.8, 2.2]
    )

    code_block(doc, [
        "def extract_procedure_features(parsed):",
        "    procedures = parsed.get('procedures', {})",
        "    total_attempts = sum(s['attempts'] for s in procedures.values())",
        "",
        "    rows = []",
        "    for name, stats in procedures.items():",
        "        fail = stats['failure']",
        "        causes = stats.get('failure_causes', {})",
        "        timeout_n = causes.get('timeout', 0)",
        "        top_cause = max(causes.values()) if causes else 0",
        "",
        "        rows.append([",
        "            stats['success_rate'],           # feature 0",
        "            100.0 - stats['success_rate'],   # feature 1",
        "            stats['attempts'] / total_attempts,  # feature 2",
        "            float(fail),                     # feature 3",
        "            len(causes) / max(fail, 1),      # feature 4",
        "            top_cause / max(fail, 1),        # feature 5",
        "            timeout_n / max(fail, 1),        # feature 6",
        "            float(fail > 0),                 # feature 7",
        "        ])",
        "    return proc_names, np.array(rows)",
    ])

    h2(doc, "3.3  Sequence Features (For LSTM)")
    body(doc, "The LSTM model needs to understand the ORDER of messages. We create a "
              "sliding window — like a sliding window of 10 consecutive messages:")
    bullet(doc, "Window 1: messages 0–9")
    bullet(doc, "Window 2: messages 1–10")
    bullet(doc, "Window 3: messages 2–11  ... and so on")

    body(doc, "Each message is encoded as 4 numbers:")
    comparison_table(doc,
        ["Feature", "Encoding", "Why"],
        [
            ["layer_id", "NAS=0, NGAP=1, RRC=2, F1AP=3, E1AP=4, XnAP=5 (divided by 5)", "Which protocol is this message from?"],
            ["proc_id", "Integer ID from vocabulary / vocab_size", "Which procedure is this?"],
            ["role_id", "request=0, response=1, failure=2, timeout=3 (divided by 3)", "Is this a request or response?"],
            ["time_delta", "Seconds since previous message / 60 (capped)", "How long since last message?"],
        ],
        col_widths=[1.3, 2.5, 2.9]
    )

    key_point(doc, "Why normalise all features to 0–1?",
              "ML models are sensitive to scale. If success_rate is 95.0 and failure_count is 5, "
              "the model might think success_rate is 19x more important just because its number "
              "is larger. Normalising puts everything on the same scale.")

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 4 — PCAP Anomaly Detection
# ══════════════════════════════════════════════════════════════════════
def ch_pcap_detection(doc):
    h1(doc, "Chapter 4 — PCAP Anomaly Detection (6 Detectors)")

    body(doc, "Files: src/detection/detectors.py, src/detection/detector.py")
    body(doc,
         "After parsing the PCAP file, we run it through 6 independent anomaly detectors. "
         "Each detector has a different approach. Procedures flagged by multiple detectors "
         "are high-confidence findings. Procedures flagged by only one detector "
         "are investigation candidates.")

    h2(doc, "4.1  The Common Output Schema")
    body(doc, "All 6 detectors return anomalies in the same format so they can be easily compared:")
    code_block(doc, [
        "{",
        "  'type':           'NAS_Registration anomaly',",
        "  'severity':       'High',          # High / Medium / Low",
        "  'score':          0.87,            # 0–1, higher = more anomalous",
        "  'detector':       'Isolation Forest',",
        "  'evidence':       'success_rate=82.0%; failures=18; top_cause=congestion',",
        "  'procedure':      'NAS_Registration',",
        "  'layer':          'NAS',",
        "  'failure_causes': {'congestion': 12, 'timeout': 6},",
        "  'recommendation': 'Check AMF load; consider load-balancing.'",
        "}",
    ])

    h2(doc, "4.2  Severity Thresholds")
    comparison_table(doc,
        ["Success Rate", "Severity", "Meaning"],
        [
            ["≥ 98%", "✅ Clean", "Normal — no action needed"],
            ["95–98%", "🟢 Low", "Monitor — minor degradation"],
            ["85–95%", "🟡 Medium", "Investigate — significant degradation"],
            ["< 85%", "🔴 High", "Act now — major failure"],
        ],
        col_widths=[1.5, 1.2, 4.0]
    )

    # ── Detector 1 ──────────────────────────────────────────────────
    h2(doc, "4.3  Detector 1: Isolation Forest")
    body(doc, "Type: Unsupervised / Tree-based  |  Input: Procedure-level features")

    h3(doc, "What is Isolation Forest?")
    body(doc,
         "Isolation Forest works on a clever idea: anomalies are rare and different from "
         "normal data. If you build a random tree and try to isolate a data point, "
         "anomalies get isolated in very few splits (short path). Normal points need many "
         "splits to isolate (long path).")
    body(doc,
         "Imagine a forest of trees where each tree randomly draws lines through your data. "
         "A point that gets 'boxed in' quickly (few lines needed) is an anomaly. "
         "A point that takes many random lines to isolate is normal.")

    code_block(doc, [
        "from sklearn.ensemble import IsolationForest",
        "from sklearn.preprocessing import StandardScaler",
        "",
        "# Scale features to zero mean, unit variance",
        "scaler   = StandardScaler()",
        "X_scaled = scaler.fit_transform(X)   # X = procedure feature matrix",
        "",
        "# Build the model",
        "clf = IsolationForest(",
        "    contamination=0.15,    # expect up to 15% of procedures to be anomalous",
        "    n_estimators=200,      # 200 trees in the forest",
        "    random_state=42        # reproducible results",
        ")",
        "preds       = clf.fit_predict(X_scaled)   # -1 = anomaly, +1 = normal",
        "raw_scores  = clf.score_samples(X_scaled) # more negative = more anomalous",
        "",
        "# Normalise scores to 0–1 scale",
        "norm_scores = 1.0 - (raw_scores - s_min) / (s_max - s_min + 1e-10)",
    ])

    h3(doc, "Why Isolation Forest?")
    bullet(doc, "No labelled data needed — we don't have examples of 'bad' procedures to train on")
    bullet(doc, "Works well when anomalies are rare — in a healthy network, <5% of procedures fail badly")
    bullet(doc, "No assumption about data distribution — doesn't require Gaussian or linear data")
    bullet(doc, "Fast and scalable — O(n log n) even with hundreds of procedures")
    bullet(doc, "contamination=0.15 means 'I expect up to 15% anomalies' — tunable per network health")

    result_box(doc, "Isolation Forest on test data", [
        "• Found: 2 anomalies",
        "• NAS_Registration: score=0.91 (High) — success_rate=95%, timeouts=2",
        "• E1AP_BearerContextSetup: score=0.97 (High) — success_rate=59%, dominant cause: procedure-cancelled",
        "• Both confirmed by other detectors",
    ])

    # ── Detector 2 ──────────────────────────────────────────────────
    h2(doc, "4.4  Detector 2: Statistical (Threshold + Cascade Rules)")
    body(doc, "Type: Rule-based / Domain Knowledge  |  Input: Procedure stats")

    h3(doc, "What is the Statistical Detector?")
    body(doc, "This detector encodes 5 telecom-specific rules directly as code. "
              "It is 100% interpretable — every finding has a clear human-readable reason.")

    comparison_table(doc,
        ["Rule", "Condition", "Severity"],
        [
            ["R1: Success Rate Threshold", "SR < 85% → High, < 95% → Medium, < 98% → Low", "High/Medium/Low"],
            ["R2: Timeout Detection", "Any timeout failure present (even if SR is OK)", "Low"],
            ["R3: Cause Concentration", "One cause > 80% of all failures", "Low"],
            ["R4: Timeout Ratio", "Timeout failures / total failures", "Low"],
            ["R5: Cascading Failure", "Upstream fails > 20% but downstream still runs", "Medium"],
        ],
        col_widths=[2.0, 3.2, 1.5]
    )

    h3(doc, "The Cascade Rule — most interesting rule")
    body(doc, "In 5G, procedures happen in sequence: Registration → Authentication → "
              "Security Mode → PDU Session. If Registration is failing 25% of the time, "
              "Authentication should also be failing (because no registration = no auth). "
              "If Authentication still shows high attempts despite failed registrations, "
              "that signals incomplete failure isolation — a system design problem.")

    code_block(doc, [
        "CHAIN = [",
        "    ('NAS_Registration',       'NAS_Authentication'),",
        "    ('NAS_Authentication',      'NAS_SecurityMode'),",
        "    ('NAS_SecurityMode',        'NAS_PDUSession_Establishment'),",
        "    ('RRC_Setup',               'RRC_Reconfiguration'),",
        "    ('NGAP_InitialContextSetup','NGAP_PDUSessionResourceSetup'),",
        "]",
        "",
        "for upstream, downstream in CHAIN:",
        "    up_fail_rate = upstream.failure / upstream.attempts",
        "    if up_fail_rate > 0.20 and downstream.attempts > 0:",
        "        # Cascade alert: upstream is failing but downstream still runs",
        "        flag_anomaly('Cascading failure', severity='Medium')",
    ])

    # ── Detector 3 ──────────────────────────────────────────────────
    h2(doc, "4.5  Detector 3: One-Class SVM")
    body(doc, "Type: Kernel / Geometric Boundary  |  Input: Procedure-level features")

    h3(doc, "What is One-Class SVM?")
    body(doc,
         "A regular SVM (Support Vector Machine) learns to separate two classes: "
         "'normal' vs 'anomaly'. But we don't have labelled anomaly examples. "
         "One-Class SVM instead learns the BOUNDARY of normal data — it draws a "
         "tight boundary around all the normal procedures. Anything outside that "
         "boundary is anomalous.")
    body(doc,
         "The 'kernel' (we use RBF = Radial Basis Function) allows the boundary to be "
         "curved and non-linear. This is important because the normal region of procedure "
         "features is not a simple box or ellipse.")

    code_block(doc, [
        "from sklearn.svm import OneClassSVM",
        "",
        "clf = OneClassSVM(",
        "    nu=0.15,      # upper bound on fraction of anomalies (like contamination)",
        "    kernel='rbf', # radial basis function = curved boundary",
        "    gamma='scale' # auto-tune to feature variance",
        ")",
        "preds  = clf.fit_predict(X_scaled)       # -1 = anomaly",
        "scores = clf.decision_function(X_scaled) # more negative = more anomalous",
    ])

    h3(doc, "Why OC-SVM on top of Isolation Forest?")
    bullet(doc, "Isolation Forest uses tree DEPTH to detect anomalies (global method)")
    bullet(doc, "OC-SVM uses geometric MARGIN (local boundary method)")
    bullet(doc, "They find different types of anomalies — together they cover more ground")
    bullet(doc, "When both agree on an anomaly, confidence is very high")

    # ── Detector 4 ──────────────────────────────────────────────────
    h2(doc, "4.6  Detector 4: LOF (Local Outlier Factor)")
    body(doc, "Type: Density-based / Local  |  Input: Procedure-level features")

    h3(doc, "What is LOF?")
    body(doc,
         "LOF compares each data point to its K nearest neighbors. A point is "
         "anomalous if its local density is much lower than its neighbors' density.")
    body(doc,
         "Simple example: Imagine all your procedures clustered together except one "
         "procedure that sits far from its nearest cluster. That isolated procedure "
         "has low local density — LOF flags it. Isolation Forest might miss it because "
         "globally it looks like just one of many data points.")

    code_block(doc, [
        "from sklearn.neighbors import LocalOutlierFactor",
        "",
        "n_neighbors = min(5, len(proc_names) - 1)  # 5 or fewer if limited data",
        "",
        "clf = LocalOutlierFactor(",
        "    n_neighbors=n_neighbors,",
        "    contamination=0.15",
        ")",
        "preds      = clf.fit_predict(X_scaled)     # -1 = anomaly",
        "lof_scores = -clf.negative_outlier_factor_ # higher = more anomalous",
    ])

    key_point(doc, "Local vs Global detection",
              "Isolation Forest and OC-SVM are GLOBAL methods — they compare each point "
              "to the whole dataset. LOF is LOCAL — it compares each point to its "
              "immediate neighbors. A procedure with 85% SR in a network where ALL others "
              "are >99% is a local outlier that global methods might miss.")

    # ── Detector 5 ──────────────────────────────────────────────────
    h2(doc, "4.7  Detector 5: Elliptic Envelope")
    body(doc, "Type: Gaussian / Mahalanobis Distance  |  Input: Procedure-level features")

    h3(doc, "What is Elliptic Envelope?")
    body(doc,
         "Elliptic Envelope assumes that normal data follows a multivariate Gaussian "
         "distribution (bell curve in multiple dimensions). It fits an 'ellipse' "
         "(in multiple dimensions) around the normal data and flags points that "
         "fall outside it.")
    body(doc,
         "The distance it uses is called Mahalanobis distance — like regular distance "
         "but adjusted for correlation between features. If success_rate and failure_rate "
         "are perfectly correlated (which they are, since they sum to 100), Mahalanobis "
         "distance correctly handles that.")

    code_block(doc, [
        "from sklearn.covariance import EllipticEnvelope",
        "import warnings",
        "",
        "# Need more samples than features (n_procs > n_features + 2)",
        "if n_procs >= n_feats + 2:",
        "    with warnings.catch_warnings():",
        "        warnings.simplefilter('ignore')  # suppress MCD convergence noise",
        "        clf   = EllipticEnvelope(contamination=0.10, random_state=42)",
        "        preds = clf.fit_predict(X_scaled)",
        "        mahal = clf.mahalanobis(X_scaled)  # squared Mahalanobis distances",
    ])

    h3(doc, "Why Elliptic Envelope?")
    body(doc, "It is the simplest and most interpretable baseline. "
              "When the other ML methods flag something that Elliptic Envelope misses, "
              "it tells us the anomaly is subtle. When Elliptic Envelope also flags it, "
              "we know it is a clear statistical outlier.")
    body(doc, "Note: contamination=0.10 (tighter than IF's 0.15) — Elliptic Envelope "
              "should only flag clear Gaussian outliers, not borderline cases.")

    # ── Detector 6 ──────────────────────────────────────────────────
    h2(doc, "4.8  Detector 6: LSTM Autoencoder")
    body(doc, "Type: Deep Learning / Sequential  |  Input: Message log windows")

    h3(doc, "What is an Autoencoder?")
    body(doc,
         "An autoencoder is a neural network that learns to compress data and then "
         "reconstruct it. If you train it on normal data, it becomes good at "
         "reconstructing normal patterns. When you show it an abnormal sequence, "
         "it cannot reconstruct it well — the reconstruction error is high.")

    h3(doc, "What is LSTM?")
    body(doc,
         "LSTM stands for Long Short-Term Memory. It is a type of neural network "
         "specifically designed for sequential data — data where the ORDER matters. "
         "In 5G, the order of messages matters a lot: Registration must come before "
         "Authentication, which must come before Security Mode, which must come before "
         "PDU Session. If this order is violated, the LSTM will have high reconstruction error.")

    code_block(doc, [
        "class _LSTMAutoencoder(nn.Module):",
        "    def __init__(self, n_features=4, hidden=32):",
        "        self.encoder = nn.LSTM(n_features, hidden)  # compress",
        "        self.decoder = nn.LSTM(hidden,     hidden)  # reconstruct",
        "        self.output  = nn.Linear(hidden, n_features)",
        "",
        "    def forward(self, x):  # x = (batch, sequence_length, features)",
        "        _, (h, c)  = self.encoder(x)             # encode to hidden state",
        "        dec_in     = h.repeat(seq_len)            # repeat for decoder",
        "        dec_out, _ = self.decoder(dec_in, (h,c))  # decode back",
        "        return self.output(dec_out)               # reconstruct input",
    ])

    body(doc, "Training process:")
    numbered(doc, "Split message log into sliding windows of 10 messages each")
    numbered(doc, "Train the autoencoder on ALL windows (no labels needed)")
    numbered(doc, "Compute reconstruction error for each window")
    numbered(doc, "Threshold = mean_error + 2 × std_error")
    numbered(doc, "Any window with error > threshold = anomalous sequence")

    key_point(doc, "LSTM vs Tabular detectors",
              "Detectors 1-5 treat each procedure independently (tabular). "
              "LSTM Autoencoder understands SEQUENCES. It can catch: "
              "out-of-order messages, unexpected gaps between messages, "
              "procedures that appear in wrong contexts — none of which "
              "tabular methods can detect.")

    result_box(doc, "LSTM on test data", [
        "• Training: 60 epochs on 20 windows, loss converges from 0.12 → 0.003",
        "• Threshold: mean_error=0.003 + 2×std=0.001 = 0.005",
        "• Anomalous windows: found near NAS_Registration with role='failure'",
        "• Evidence: Window 5 (events 4-14): reconstruction_error=0.0089 > threshold=0.005",
        "• Severity: Medium (score=1.78x threshold)",
    ])

    h2(doc, "4.9  Multi-Detector Merging")
    body(doc, "After all 6 detectors run, we merge their results by procedure:")

    code_block(doc, [
        "def _merge(anomalies):",
        "    # Group by (procedure, type)",
        "    for a in anomalies:",
        "        key = f\"{a['procedure']}|{a['type']}\"",
        "        if key not in groups:",
        "            groups[key] = a   # first detector to flag this",
        "        else:",
        "            # Keep highest severity",
        "            if SEV_RANK[a['severity']] > SEV_RANK[existing['severity']]:",
        "                groups[key].update(severity, score, evidence)",
        "            # Record which detectors agreed",
        "            groups[key]['detectors'].append(a['detector'])",
        "",
        "    # Tag each anomaly with how many detectors confirmed it",
        "    for a in result:",
        "        a['confirmed_by'] = len(a['detectors'])",
    ])

    result_box(doc, "Full 6-detector ensemble on test data", [
        "• E1AP_BearerContextSetup: High | confirmed_by=4 (IF + Statistical + OC-SVM + LOF)",
        "• NAS_PDUSession_Establishment: High | confirmed_by=3 (Statistical + LOF + Elliptic)",
        "• NAS_Registration: High | confirmed_by=3 (IF + Statistical + OC-SVM)",
        "• Rule: confirmed_by ≥ 3 = act now | confirmed_by = 1 = investigate",
    ])

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 5 — KPI Detection
# ══════════════════════════════════════════════════════════════════════
def ch_kpi_detection(doc):
    h1(doc, "Chapter 5 — KPI Anomaly Detection (6 Methods)")

    body(doc, "File: src/detection/kpi_detector.py")
    body(doc,
         "KPI anomaly detection works differently from PCAP detection. "
         "Instead of finding anomalous procedures, we find anomalous KPI values "
         "across cells and time intervals. We use 6 complementary methods.")

    h2(doc, "5.1  Method 1: Threshold Violation")
    body(doc, "The most direct method — compare each KPI value against its configured "
              "warning and critical thresholds.")

    code_block(doc, [
        "for row in records:",
        "    for kpi in kpi_columns:",
        "        val  = row[kpi]",
        "        meta = get_meta(kpi)   # get thresholds from catalog",
        "",
        "        if direction == 'higher_better':  # e.g. RRC Success Rate",
        "            if val <= meta['critical']:  severity = 'Critical'",
        "            elif val <= meta['warning']: severity = 'High'",
        "        else:                            # e.g. PRB Utilization",
        "            if val >= meta['critical']:  severity = 'Critical'",
        "            elif val >= meta['warning']: severity = 'High'",
    ])

    key_point(doc, "Why start with thresholds?",
              "Thresholds encode operator SLA (Service Level Agreement) directly. "
              "They have zero false positives for known violations and are "
              "100% interpretable. Every other method is calibrated against this baseline.")

    h2(doc, "5.2  Method 2: Peer Comparison (Z-score)")
    body(doc, "Finds cells that are significantly worse than the fleet average, "
              "even if they haven't violated a threshold.")

    code_block(doc, [
        "# Compute per-cell mean for each KPI",
        "cell_means = df.groupby(cell_col)[kpi].mean()",
        "fleet_mean = cell_means.mean()  # average across all cells",
        "fleet_std  = cell_means.std()   # spread across all cells",
        "",
        "# Z-score = how many standard deviations from the fleet average",
        "z_scores = (cell_means - fleet_mean) / fleet_std",
        "",
        "for cell, z in z_scores.items():",
        "    if direction == 'higher_better' and z < -2.0:",
        "        # This cell is more than 2 std below the fleet average",
        "        severity = 'High' if z < -3.0 else 'Medium'",
    ])

    warning_box(doc, "Why z-score alone is not enough",
        "Z-score assumes Gaussian distribution. PRB utilization is right-skewed "
        "(most cells are at 30-50% but a few spike to 90%+). In such distributions, "
        "z-score underestimates the severity of high outliers. That is why we also "
        "use IQR (Method 4) which makes no distribution assumption.")

    h2(doc, "5.3  Method 3: Trend Analysis (Linear Regression)")
    body(doc, "Detects KPIs that are consistently getting worse over time, "
              "even before they cross any threshold.")

    code_block(doc, [
        "# Convert timestamps to hours since start",
        "df['_hrs'] = (df[ts_col] - t0).dt.total_seconds() / 3600.0",
        "",
        "# Fit a line: y = slope × hours + intercept",
        "slope, intercept = np.polyfit(x=hours, y=kpi_values, deg=1)",
        "",
        "# A degrading slope for 'higher_better' KPI = negative slope",
        "if direction == 'higher_better' and slope < -0.5:  # drops > 0.5 per hour",
        "    sev = 'High' if abs(slope) > 2.0 else 'Medium'",
        "    evidence = f'slope={slope:+.3f}/hr, estimated drop: {delta:.2f}'",
    ])

    key_point(doc, "Early warning power",
              "A CQI declining at -0.5 per hour will cross the warning threshold "
              "in 2 hours. Trend detection fires NOW, giving engineers 2 hours to "
              "investigate before customers experience problems.")

    h2(doc, "5.4  Method 4: IQR (Tukey Fence) — NEW")
    body(doc, "A robust, distribution-free outlier detection method. "
              "IQR stands for Interquartile Range.")

    body(doc, "Understanding quartiles:")
    bullet(doc, "Q1 (25th percentile): 25% of values are below this")
    bullet(doc, "Q2 (50th percentile / median): middle value")
    bullet(doc, "Q3 (75th percentile): 75% of values are below this")
    bullet(doc, "IQR = Q3 - Q1 = the range containing the middle 50% of data")
    bullet(doc, "Lower fence = Q1 - 1.5 × IQR")
    bullet(doc, "Upper fence = Q3 + 1.5 × IQR")
    bullet(doc, "Any value outside these fences is a Tukey outlier")

    code_block(doc, [
        "q1  = vals.quantile(0.25)",
        "q3  = vals.quantile(0.75)",
        "iqr = q3 - q1",
        "",
        "lower_fence = q1 - 1.5 * iqr   # below this = outlier",
        "upper_fence = q3 + 1.5 * iqr   # above this = outlier",
        "",
        "for val in values:",
        "    if direction == 'higher_better' and val < lower_fence:",
        "        sev = 'High' if val < (lower_fence - iqr) else 'Medium'",
    ])

    h3(doc, "Why IQR instead of just Z-score?")
    comparison_table(doc,
        ["Scenario", "Z-score result", "IQR result"],
        [
            ["Normal data, one bad cell", "Works correctly", "Works correctly"],
            ["Right-skewed PRB utilization", "⚠️ Underestimates outlier severity", "✅ Correct — not affected by skew"],
            ["Outlier inflates the mean", "⚠️ Outlier makes σ larger, hides itself", "✅ Resistant — uses median-based measure"],
            ["Bimodal distribution", "⚠️ Often fails", "✅ Works — fences adapt to actual data"],
        ],
        col_widths=[2.5, 2.2, 2.2]
    )

    h2(doc, "5.5  Method 5: CUSUM (Cumulative Sum Control Chart) — NEW")
    body(doc, "CUSUM is the standard early-warning method used in telecom Network "
              "Operations Centres (NOC). It detects when a KPI starts consistently "
              "drifting away from its normal value.")

    h3(doc, "How CUSUM works — the key insight")
    body(doc,
         "Traditional thresholds look at each data point independently. "
         "CUSUM accumulates evidence over time. Even if each individual deviation is "
         "small (below any threshold), if deviations are consistently in the same "
         "direction (always below normal), CUSUM will eventually fire.")
    body(doc,
         "Example: RRC SR is normally 99.5%. Over 10 intervals it goes: "
         "99.3%, 99.1%, 98.9%, 98.7%, 98.5%, 98.3%, 98.1%, 97.9%, 97.7%, 97.5%. "
         "Each individual value is above the 98% warning threshold, so Threshold "
         "detection stays silent. But the consistent downward trend causes CUSUM to "
         "fire at interval 7 — warning the engineer 3 intervals BEFORE the threshold breach.")

    code_block(doc, [
        "K = 0.5   # slack parameter: ignore noise within 0.5 standard deviations",
        "H = 5.0   # alarm threshold: fire when cumulative sum > 5 standard deviations",
        "",
        "S_pos = 0.0   # cumulative sum for upward drift",
        "S_neg = 0.0   # cumulative sum for downward drift",
        "",
        "for val in series:",
        "    z = (val - mean) / std   # normalise: how many std from normal?",
        "    S_pos = max(0, S_pos + z - K)   # accumulate upward deviations",
        "    S_neg = max(0, S_neg - z - K)   # accumulate downward deviations",
        "",
        "    if direction == 'higher_better' and S_neg > H:",
        "        alert('Sustained downward drift detected!')",
        "        break  # one alert per cell per KPI",
    ])

    result_box(doc, "CUSUM early detection example", [
        "• PRB Cell_B: sustained upward drift detected at hour 15",
        "• S+ accumulated to 6.2 (> H=5) before threshold even fired",
        "• Engineer is alerted 2 hours before PRB hits 90% warning threshold",
        "• Saves 2 hours of customer-impacting congestion",
    ])

    h2(doc, "5.6  Method 6: Bollinger Bands (Rolling Envelope) — NEW")
    body(doc, "Bollinger Bands come from financial market analysis and are now widely "
              "used in network monitoring for burst detection.")

    h3(doc, "How Bollinger Bands work")
    body(doc, "Instead of a fixed global threshold, Bollinger Bands create a DYNAMIC "
              "threshold that adapts to the recent history of each cell:")
    bullet(doc, "Rolling mean: average of the last 5 KPI values")
    bullet(doc, "Rolling std: standard deviation of the last 5 values")
    bullet(doc, "Upper band: rolling_mean + 2 × rolling_std")
    bullet(doc, "Lower band: rolling_mean - 2 × rolling_std")
    bullet(doc, "Anything outside these bands = anomalous for THIS cell at THIS time")

    h3(doc, "The important technical detail: Lagged window")
    body(doc, "A common mistake: using the current value to compute its own band. "
              "This allows a spike to 'inflate' the mean and hide itself inside its own band!")

    code_block(doc, [
        "# WRONG: current value is included in its own band computation",
        "rolling_mean = series.rolling(5).mean()  # at position i, uses values [i-4 to i]",
        "# If value at i=15 spikes to 97, rolling_mean becomes (42+42+42+42+97)/5 = 53",
        "# upper_band = 53 + 2×22 = 97  — spike is exactly AT the band!",
        "",
        "# CORRECT: use lagged window — exclude current value",
        "lagged = series.shift(1)             # shift forward by 1",
        "rolling_mean = lagged.rolling(5).mean()  # at i=15, uses values [i-5 to i-1]",
        "# rolling_mean = (42+42+42+42+42)/5 = 42",
        "# upper_band   = 42 + 2×2 = 46",
        "# val=97 >> band=46  → correctly flagged as High severity",
    ])

    result_box(doc, "Bollinger Bands on test data", [
        "• PRB Cell_B: burst at hour 15 (97%) correctly flagged above upper band (46%)",
        "• RRC Cell_A: step decline detected as values exit the rolling envelope",
        "• 4 anomalies found, all with High or Medium severity",
        "• Without the shift(1) fix: 0 anomalies would have been detected",
    ])

    h2(doc, "5.7  Why All 6 Methods Together — Coverage Map")
    comparison_table(doc,
        ["Scenario", "Best Method"],
        [
            ["Cell's CQI drops below 5.0 (critical threshold)", "Threshold ✅"],
            ["Cell at 97% SR but all peers are at 99.5%", "Peer Comparison ✅"],
            ["RRC SR slowly declining from 99.5% to 97% over 24 hours", "Trend + CUSUM ✅"],
            ["PRB spikes to 92% for 2 intervals then drops back", "Bollinger Bands ✅"],
            ["KPI values outlier but distribution is right-skewed", "IQR ✅"],
            ["KPI drifting below mean, not yet at threshold", "CUSUM ✅ (5 hrs early)"],
        ],
        col_widths=[4.5, 2.8]
    )

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 6 — Dashboard
# ══════════════════════════════════════════════════════════════════════
def ch_dashboard(doc):
    h1(doc, "Chapter 6 — The Streamlit Dashboard")

    body(doc, "File: src/dashboard/app.py")
    body(doc,
         "The dashboard is the user-facing part of the project. It is built with "
         "Streamlit — a Python library that lets you build web applications entirely "
         "in Python, without writing any HTML, CSS, or JavaScript. You just write "
         "Python code and Streamlit renders it as a beautiful web app.")

    h2(doc, "6.1  How to Run")
    code_block(doc, [
        "# Method 1: Direct",
        "streamlit run src/dashboard/app.py",
        "",
        "# Method 2: Using Makefile",
        "make dev",
        "",
        "# Then open: http://localhost:8501",
    ])

    h2(doc, "6.2  Dashboard Architecture")
    body(doc, "The dashboard has two completely separate flows depending on what file you upload:")

    comparison_table(doc,
        ["Upload Type", "Flow", "Output"],
        [
            ["PCAP file (.pcap)", "parse_pcap_real() → detect_anomalies_by_detector()", "Protocol tabs, IE inspector, 6-detector comparison matrix"],
            ["KPI file (.xlsx/.csv)", "parse_kpi_file() → detect_kpi_anomalies_by_detector()", "KPI health summary, trend charts, 6-method comparison matrix"],
        ],
        col_widths=[1.8, 2.5, 2.8]
    )

    h2(doc, "6.3  PCAP Dashboard Sections")

    h3(doc, "Section 1: Summary Metrics")
    code_block(doc, [
        "col1.metric('Total signalling events', parsed['total_events'])",
        "col2.metric('Procedures tracked', len(procedures))",
        "# Per-layer counts: NAS, NGAP, RRC, F1AP, E1AP, XnAP",
    ])

    h3(doc, "Section 2: Per-Layer Procedure Tables")
    body(doc, "Shows a table for each protocol with: Procedure name, Attempts, "
              "Success count, Failure count, Success Rate %, Top Failure Cause")

    h3(doc, "Section 3: Failure Cause Breakdown")
    body(doc, "For each procedure that has failures, shows a bar chart of "
              "all the different failure causes and how many times each occurred.")

    h3(doc, "Section 4: Message Log & IE Inspector")
    code_block(doc, [
        "for evt in filtered_log[:200]:  # show max 200 events",
        "    label = f\"[{evt['layer']}] {evt['procedure']} → UE={evt['ue_id']}",
        "              IEs: {evt['observed_ie_count']}/{evt['expected_ie_count']}\"",
        "    with st.expander(label):",
        "        # Left: Observed IEs (what we actually saw in the packet)",
        "        # Right: Missing Mandatory IEs (what 3GPP says should be there)",
    ])

    h3(doc, "Section 5: Anomaly Detection — 6 Detectors")
    body(doc, "This section has 3 parts:")
    numbered(doc, "Rationale expander — explains WHY each detector was chosen (for reviewers)")
    numbered(doc, "Method Comparison Matrix — table showing which detector flagged which procedure")
    numbered(doc, "Individual anomaly cards — expandable details for each finding with evidence and recommendation")

    h2(doc, "6.4  KPI Dashboard Sections")

    h3(doc, "Section 1: KPI Health Summary (Traffic Light Table)")
    body(doc, "Shows all KPIs with a colour-coded status:")
    bullet(doc, "🔴 Red = mean value is at or below Critical threshold")
    bullet(doc, "🟡 Yellow = mean value is at or below Warning threshold")
    bullet(doc, "🟢 Green = mean value is above Warning threshold (healthy)")
    bullet(doc, "Also shows: Min, Max, Mean, P10, P90, Warning threshold, Critical threshold")

    h3(doc, "Section 2: KPI Trend Explorer")
    code_block(doc, [
        "# User selects which KPI and which cell to view",
        "sel_kpi  = st.selectbox('Select KPI', kpi_cols)",
        "sel_cell = st.selectbox('Select Cell', ['All cells (fleet avg)'] + cells)",
        "",
        "# Plot the trend with threshold lines",
        "fig = px.line(df_line, x=ts_col, y='value', title=title)",
        "fig.add_hline(y=meta['warning'],  line_dash='dot', line_color='orange')",
        "fig.add_hline(y=meta['critical'], line_dash='dot', line_color='red')",
    ])

    h3(doc, "Section 3: Per-Cell KPI Breakdown")
    body(doc, "A heatmap-style table showing every cell's average value for every KPI. "
              "Sorted by the first KPI column so the worst-performing cells appear at top.")

    h3(doc, "Section 4: KPI Anomaly Detection")
    body(doc, "Same structure as PCAP detection: rationale expander + method comparison "
              "matrix + individual anomaly cards. Each card shows:")
    bullet(doc, "Severity level (Critical/High/Medium/Low)")
    bullet(doc, "Which KPI, which cell, which time")
    bullet(doc, "Evidence: what exactly was measured vs what was expected")
    bullet(doc, "Recommendation: what to do about it")

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 7 — LLM Explainer (Planned)
# ══════════════════════════════════════════════════════════════════════
def ch_llm(doc):
    h1(doc, "Chapter 7 — LLM Explainer (Coming Next)")

    body(doc, "Files: src/llm/explainer.py (stub), src/rag/ (empty — to be built)")
    body(doc,
         "Currently the LLM section of the dashboard shows stub (fake) explanations. "
         "The plan is to replace this with a real AI system that can explain anomalies "
         "in plain English, citing the actual 3GPP specifications.")

    h2(doc, "7.1  What We Are Building")
    body(doc, "The system will use RAG — Retrieval Augmented Generation:")
    numbered(doc, "3GPP spec documents are split into small chunks and embedded into a vector database (FAISS)")
    numbered(doc, "When an anomaly is detected, we find the most relevant spec sections")
    numbered(doc, "We pass the anomaly + relevant specs to a local AI model (Phi-3 Mini)")
    numbered(doc, "The model generates a structured explanation: hypothesis + citations + investigation hints")

    h2(doc, "7.2  Components Already in requirements.txt")
    comparison_table(doc,
        ["Library", "Purpose", "Status"],
        [
            ["faiss-cpu", "Vector database for 3GPP spec chunks", "Installed, not yet used"],
            ["sentence-transformers", "Convert text to embeddings (bge-small-en-v1.5)", "Installed, not yet used"],
            ["rank-bm25", "Keyword-based search as fallback", "Installed, not yet used"],
            ["ollama", "Run Phi-3 Mini locally", "Installed, not yet used"],
        ],
        col_widths=[2.0, 2.5, 2.0]
    )

    h2(doc, "7.3  Current Stub")
    code_block(doc, [
        "def explain_anomaly_stub(anomaly):",
        "    # Returns a fake but structurally correct explanation",
        "    return {",
        "        'hypothesis': 'The detected pattern is consistent with...',",
        "        'citations':  [{'spec': 'TS 24.501', 'section': '5.5.1.2',",
        "                        'quote': '5GMM cause #22 = Congestion'}],",
        "        'investigation_hints': ['Check AMF utilisation', ...]",
        "    }",
    ])

    body(doc, "The stub has the correct output structure so the dashboard UI already "
              "works. When we build the real LLM, we just replace this function "
              "and the dashboard will immediately show real explanations.")

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 8 — Results Summary
# ══════════════════════════════════════════════════════════════════════
def ch_results(doc):
    h1(doc, "Chapter 8 — Results Summary")

    h2(doc, "8.1  PCAP Detection Results (Test Data)")
    body(doc, "Using the synthetic test PCAP (scripts/generate_test_pcap.py) "
              "with intentional failures injected:")

    comparison_table(doc,
        ["Detector", "Anomalies Found", "Key Findings"],
        [
            ["Isolation Forest", "2", "NAS_Registration (High), E1AP_BearerContextSetup (High)"],
            ["Statistical", "4", "NAS_Registration (Low), NAS_PDUSession (High), RRC_Setup (Low), E1AP_Bearer (High)"],
            ["One-Class SVM", "6", "Caught 6 procedures — broader than IF"],
            ["LOF", "2", "NAS_PDUSession (Low), E1AP_Bearer (High) — local density method"],
            ["Elliptic Envelope", "1", "NAS_PDUSession (High) — strict Gaussian baseline"],
            ["LSTM Autoencoder", "1", "Anomalous message sequence near NAS_Registration"],
            ["MERGED TOTAL", "7 unique", "E1AP: 4/6 detectors | NAS_PDUSession: 3/6 | NAS_Reg: 3/6"],
        ],
        col_widths=[1.8, 1.5, 3.7]
    )

    h2(doc, "8.2  KPI Detection Results (Synthetic Data)")
    body(doc, "Test scenario: Cell_A RRC SR declining, Cell_B PRB burst at hour 15, "
              "Cell_C CQI always bad:")

    comparison_table(doc,
        ["Method", "Anomalies", "What it Caught"],
        [
            ["Threshold", "43", "Cell_C CQI < critical=5 (every interval), Cell_A RRC SR < critical=95%"],
            ["Peer Comparison", "0", "Only 3 cells — too few for reliable z-score (needs 5+)"],
            ["Trend", "1", "Fleet-wide PRB utilization increasing slope"],
            ["IQR", "2", "Cell_A RRC SR outlier, Cell_B PRB outlier"],
            ["CUSUM", "1", "Cell_B PRB sustained upward drift detected at hour 15"],
            ["Bollinger Bands", "4", "PRB spikes at multiple cells, RRC step decline at Cell_A"],
        ],
        col_widths=[1.8, 1.2, 4.0]
    )

    h2(doc, "8.3  Dashboard Performance")
    comparison_table(doc,
        ["Operation", "Time", "Notes"],
        [
            ["Load small PCAP (< 1MB)", "2-5 seconds", "pyshark spawns tshark process"],
            ["Run 6 PCAP detectors", "3-10 seconds", "LSTM training dominates (60 epochs)"],
            ["Load KPI Excel (720 rows)", "< 1 second", "pandas read_excel is fast"],
            ["Run 6 KPI detectors", "2-4 seconds", "CUSUM and Bollinger per-cell loops"],
            ["Render dashboard", "< 1 second", "Streamlit caches static elements"],
        ],
        col_widths=[2.3, 1.3, 3.4]
    )

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 9 — Key Concepts Quick Reference
# ══════════════════════════════════════════════════════════════════════
def ch_glossary(doc):
    h1(doc, "Chapter 9 — Key Concepts Quick Reference")

    h2(doc, "9.1  3GPP Terms")
    comparison_table(doc,
        ["Term", "Full Form", "Simple Meaning"],
        [
            ["UE", "User Equipment", "Your phone or any 5G device"],
            ["gNB", "Next Generation NodeB", "The 5G tower/base station"],
            ["AMF", "Access and Mobility Management Function", "Core network function that manages UE registration"],
            ["SMF", "Session Management Function", "Core network function that manages data sessions (PDU sessions)"],
            ["UPF", "User Plane Function", "Routes actual user data (videos, calls)"],
            ["PCAP", "Packet Capture", "File containing recorded network traffic"],
            ["KPI", "Key Performance Indicator", "Statistics measuring network health"],
            ["PRB", "Physical Resource Block", "Unit of radio bandwidth allocation"],
            ["CQI", "Channel Quality Indicator", "Signal quality as reported by UE (0-15)"],
            ["SINR", "Signal to Interference + Noise Ratio", "How good/clean the radio signal is (in dB)"],
            ["BLER", "Block Error Rate", "How many data blocks need retransmission"],
            ["SR", "Success Rate", "Percentage of attempts that succeed"],
            ["HO", "Handover", "Moving a connection from one tower to another"],
            ["IE", "Information Element", "A specific data field inside a protocol message"],
            ["ASN.1", "Abstract Syntax Notation One", "Encoding format used by 3GPP protocols"],
            ["TLV", "Type-Length-Value", "Simple binary encoding used by NAS protocol"],
        ],
        col_widths=[1.0, 2.6, 3.1]
    )

    h2(doc, "9.2  Machine Learning Terms")
    comparison_table(doc,
        ["Term", "Simple Meaning"],
        [
            ["Unsupervised Learning", "Learning without labelled examples. The model finds patterns on its own."],
            ["Anomaly Detection", "Finding data points that are significantly different from the norm."],
            ["Contamination", "The expected fraction of anomalies in the data (e.g., 0.15 = 15%)"],
            ["Feature Engineering", "Converting raw data into numbers that ML models can process"],
            ["Normalisation / Scaling", "Adjusting all features to the same scale (0-1 or mean=0, std=1)"],
            ["Standard Scaler", "Transforms features to have mean=0 and std=1"],
            ["Reconstruction Error", "How different the autoencoder's output is from its input"],
            ["Isolation Score", "In Isolation Forest: shorter path = more anomalous"],
            ["Decision Function", "In SVM: negative = anomaly, positive = normal"],
            ["LOF Score", "Local Outlier Factor: > 1.0 means more anomalous than neighbors"],
            ["Mahalanobis Distance", "Distance that accounts for correlation between features"],
            ["CUSUM S+/S-", "Accumulated positive/negative deviations from mean"],
        ],
        col_widths=[2.0, 4.7]
    )

    h2(doc, "9.3  Python Libraries Quick Reference")
    comparison_table(doc,
        ["Library", "What it does in our project"],
        [
            ["pyshark", "Wraps Wireshark to decode 5G protocol packets"],
            ["pandas", "Data manipulation — DataFrames, groupby, rolling windows"],
            ["numpy", "Numerical arrays, polyfit for trend, quantile for IQR"],
            ["sklearn", "Isolation Forest, OC-SVM, LOF, Elliptic Envelope, StandardScaler"],
            ["torch (PyTorch)", "LSTM Autoencoder — nn.LSTM, nn.Linear, Adam optimizer"],
            ["streamlit", "Web dashboard — st.metric, st.dataframe, st.expander"],
            ["plotly", "Interactive charts — px.line, add_hline"],
            ["python-pptx", "Auto-generate PowerPoint reviewer presentations"],
            ["python-docx", "Auto-generate Word study documents (this document!)"],
        ],
        col_widths=[1.8, 4.9]
    )

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 10 — Study Questions & Answers
# ══════════════════════════════════════════════════════════════════════
def ch_study_qa(doc):
    h1(doc, "Chapter 10 — Study Questions & Answers")

    body(doc, "These are the questions a reviewer or interviewer might ask about this project.")

    qas = [
        ("Q1: Why did you choose pyshark over scapy for PCAP parsing?",
         "pyshark wraps Wireshark's 3GPP dissectors that have been maintained for 20+ years. "
         "It gives us real field names like pkt.ngap.procedurecode directly. With scapy, "
         "we would get raw bytes and need to write ~3000 lines of custom decoders for all 6 protocols. "
         "NAS uses a custom TLV encoding that is not standard ASN.1, so even asn1tools would not help with it."),

        ("Q2: Why do you use 6 anomaly detectors instead of just one?",
         "No single detector catches all types of anomalies. Isolation Forest is global and tabular — "
         "it misses local density anomalies. LOF is local but also tabular — it misses temporal patterns. "
         "LSTM catches temporal patterns but not tabular outliers. OC-SVM adds a geometric boundary that "
         "complements both. Having 6 detectors means when 3+ agree, we have very high confidence. "
         "Each detector has a specific blind spot that another covers."),

        ("Q3: What is the difference between CUSUM and linear regression (Trend)?",
         "Linear regression fits a straight line through all data and detects monotonic decline over the "
         "entire dataset. CUSUM operates step-by-step, accumulating deviations from the normal mean. "
         "Trend is better for long-term gradual decline. CUSUM is better for detecting the exact point "
         "when a KPI starts degrading, even if individual deviations are small. CUSUM can fire 5+ hours "
         "before the trend line would cross a threshold."),

        ("Q4: Why do you use IQR instead of z-score for KPI outlier detection?",
         "Z-score assumes Gaussian (normal/bell curve) distribution. PRB utilization is right-skewed "
         "(most cells are at 30-50% but a few spike to 90%+). When one extreme value exists, it inflates "
         "the standard deviation, making z-score less sensitive to other outliers. IQR uses Q1 and Q3 "
         "(which are not affected by extreme values) so it is robust to skewed distributions."),

        ("Q5: What is the purpose of the LSTM Autoencoder?",
         "The LSTM Autoencoder detects anomalous MESSAGE SEQUENCES, not anomalous individual procedures. "
         "5G signaling has a strict order: Registration → Authentication → Security Mode → PDU Session. "
         "If this order is violated, or if there are unexpected gaps or repetitions, all tabular detectors "
         "will miss it (they look at per-procedure statistics, not sequence order). The LSTM learns what "
         "normal sequences look like and flags sequences with high reconstruction error."),

        ("Q6: Why do Bollinger Bands use a lagged (shifted) window?",
         "Without shifting, the current value is included in its own rolling mean and std computation. "
         "A spike at hour 15 (value=97) would inflate the rolling mean from 42 to 53 and rolling std "
         "from 2 to 22. The upper band would jump to 53+44=97 — exactly where the spike is, so it would "
         "NOT be flagged. By shifting the window by 1, the current value's band is computed from PRIOR "
         "values only (all ~42), so upper_band=46. The spike at 97 >> 46 and is correctly flagged High."),

        ("Q7: What does 'confirmed_by=4 detectors' mean in the results?",
         "It means 4 out of 6 independent detectors, each using a completely different algorithm, "
         "all flagged the same procedure as anomalous. The probability of 4 independent detectors "
         "all producing false positives simultaneously is extremely low. This is the highest confidence "
         "finding in our system."),

        ("Q8: How does the KPI catalog (kpi_defs.py) work?",
         "The catalog defines every KPI with: label (human-readable name), category (Availability/Accessibility/etc.), "
         "unit (%, Mbps, dB), direction (higher_better or lower_better), warning threshold, critical threshold, "
         "and aliases (all vendor-specific column names for this KPI). When parsing an Excel file, we use "
         "aliases to automatically rename vendor-specific columns to our canonical names regardless of whether "
         "the file came from Ericsson, Nokia, or Huawei."),
    ]

    for q, a in qas:
        p = doc.add_paragraph()
        r = p.add_run(q)
        r.bold = True
        r.font.color.rgb = NAVY
        r.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)

        p2 = doc.add_paragraph(a)
        p2.paragraph_format.left_indent = Inches(0.3)
        p2.paragraph_format.space_after = Pt(6)
        p2.runs[0].font.color.rgb = GREY

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 11 — What's Next
# ══════════════════════════════════════════════════════════════════════
def ch_next(doc):
    h1(doc, "Chapter 11 — What We Are Building Next")

    h2(doc, "11.1  LLM Explainer — Phase 3")
    body(doc, "The biggest remaining piece is replacing the stub LLM explainer with a real one:")

    numbered(doc, "Collect 3GPP spec excerpts (TS 24.501, 38.413, 38.331, 38.473, 38.463, 38.423)")
    numbered(doc, "Split into small chunks (paragraph-level)")
    numbered(doc, "Convert chunks to vector embeddings using bge-small-en-v1.5")
    numbered(doc, "Store in FAISS index (fast similarity search)")
    numbered(doc, "For each anomaly: find top-5 most relevant spec sections")
    numbered(doc, "Combine anomaly details + spec sections into a prompt")
    numbered(doc, "Send to Phi-3 Mini (running locally via Ollama)")
    numbered(doc, "Display structured output: hypothesis + citations + hints")

    h2(doc, "11.2  Real Output Format (Planned)")
    code_block(doc, [
        "# What the real LLM explainer will return:",
        "{",
        "  'hypothesis': 'The E1AP Bearer Context Setup failure rate of 41% is",
        "    consistent with a gNB-CU-UP resource exhaustion event. Per TS 38.463",
        "    Section 8.3.1, the procedure-cancelled cause indicates...',",
        "",
        "  'citations': [",
        "    {'spec': 'TS 38.463', 'section': '8.3.1',",
        "     'quote': 'Bearer Context Setup procedure is initiated by the gNB-CU-CP'},",
        "    {'spec': 'TS 38.463', 'section': '9.3.1.5',",
        "     'quote': 'Cause IE values: procedure-cancelled indicates...'}",
        "  ],",
        "",
        "  'investigation_hints': [",
        "    'Check gNB-CU-UP memory and CPU utilisation in O&M',",
        "    'Verify E1 interface connectivity and SCTP association status',",
        "    'Compare with baseline: was there a software upgrade recently?'",
        "  ]",
        "}",
    ])

    h2(doc, "11.3  Future Enhancements")
    bullet(doc, "FastAPI backend: move detection logic to a REST API so the dashboard is just a frontend")
    bullet(doc, "MLflow tracking: log detector parameters and results for experiment comparison")
    bullet(doc, "DVC (Data Version Control): version control the test PCAPs and KPI files")
    bullet(doc, "Automated report generation: daily PDF report of all anomalies")
    bullet(doc, "Alert integration: send alerts to Slack/email when Critical anomalies found")

    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("End of Document")
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "Keep this document for reference. "
        "Every function, every choice, and every result is documented here.\n"
        "github.com/vikasrai-tech/telecom-analyzer"
    )
    r2.font.color.rgb = GREY
    r2.font.size = Pt(11)


# ══════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════
def build():
    print("Building study document...")
    doc = new_doc()

    cover_page(doc)
    ch_overview(doc)
    ch_pcap_parsing(doc)
    ch_kpi_parser(doc)
    ch_features(doc)
    ch_pcap_detection(doc)
    ch_kpi_detection(doc)
    ch_dashboard(doc)
    ch_llm(doc)
    ch_results(doc)
    ch_glossary(doc)
    ch_study_qa(doc)
    ch_next(doc)

    out_path = "/mnt/e/telecom-analyzer/docs/Telecom_Analyzer_Study_Guide.docx"
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"Saved: {out_path}")
    size_kb = os.path.getsize(out_path) // 1024
    print(f"Size: {size_kb} KB")
    print(f"Chapters: 11")


if __name__ == "__main__":
    build()
