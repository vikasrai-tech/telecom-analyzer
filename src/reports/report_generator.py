"""
Report generator — exports anomaly detection results to CSV, XLSX, DOCX, PDF, or HTML.

Usage:
    from src.reports.report_generator import generate_csv, generate_xlsx, generate_docx, generate_pdf, generate_html

Each function accepts:
    sections      : list of {"title": str, "df": pd.DataFrame, "notes": str (optional)}
    meta          : dict of scalar metadata shown in the report header
    figures       : (HTML only) list of {"title": str, "fig": plotly.Figure}
    anomaly_cards : (HTML only) list of {"severity", "title", "evidence",
                    "recommendation", "value", "unit", "warning", "critical"}
Returns bytes ready for st.download_button.
"""

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

import pandas as pd

ReportSection = Dict[str, Any]


# ── CSV ───────────────────────────────────────────────────────────────────────

def generate_csv(sections: List[ReportSection], meta: Dict[str, Any]) -> bytes:
    buf = io.StringIO()
    buf.write("# Telecom Analyzer — Anomaly Detection Report\n")
    buf.write(f"# Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    for k, v in meta.items():
        buf.write(f"# {k}: {v}\n")
    buf.write("\n")

    for sec in sections:
        buf.write(f"## {sec['title']}\n")
        if sec.get("notes"):
            buf.write(f"# {sec['notes']}\n")
        df: Optional[pd.DataFrame] = sec.get("df")
        if df is not None and not df.empty:
            df.to_csv(buf, index=False)
        buf.write("\n")

    return buf.getvalue().encode("utf-8")


# ── XLSX ──────────────────────────────────────────────────────────────────────

def generate_xlsx(sections: List[ReportSection], meta: Dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        # Summary sheet
        summary_rows = [{"Field": k, "Value": str(v)} for k, v in meta.items()]
        summary_rows.append(
            {"Field": "Generated", "Value": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
        )
        pd.DataFrame(summary_rows).to_excel(writer, sheet_name="Summary", index=False)

        for sec in sections:
            df: Optional[pd.DataFrame] = sec.get("df")
            if df is not None and not df.empty:
                # Excel sheet names are capped at 31 chars and must be unique
                sheet = _safe_sheet_name(sec["title"], writer.book.sheetnames)
                df.to_excel(writer, sheet_name=sheet, index=False)

    return buf.getvalue()


def _safe_sheet_name(title: str, existing: list) -> str:
    # Excel forbids: / \ ? * [ ] :  and limits to 31 chars
    import re
    clean = re.sub(r"[/\\?*\[\]:]", "-", title).strip()[:31]
    if clean not in existing:
        return clean
    for i in range(2, 100):
        candidate = re.sub(r"[/\\?*\[\]:]", "-", title).strip()[:28] + f"_{i}"
        if candidate not in existing:
            return candidate
    return clean


# ── PDF ───────────────────────────────────────────────────────────────────────

def generate_pdf(sections: List[ReportSection], meta: Dict[str, Any]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=2.0 * cm,
        bottomMargin=1.5 * cm,
    )

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("ReportTitle", parent=styles["Title"], fontSize=16, spaceAfter=6)
    h2 = ParagraphStyle(
        "SectionTitle", parent=styles["Heading2"], fontSize=11, spaceBefore=14, spaceAfter=4
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["Normal"], fontSize=9, leading=13, textColor=colors.HexColor("#444444")
    )
    note_style = ParagraphStyle(
        "Note", parent=styles["Italic"], fontSize=8, textColor=colors.HexColor("#666666"), spaceAfter=4
    )

    story = []

    # ── Header ────────────────────────────────────────────────────────────────
    story.append(Paragraph("Telecom Analyzer — Anomaly Detection Report", h1))
    story.append(
        Paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            meta_style,
        )
    )
    for k, v in meta.items():
        story.append(Paragraph(f"<b>{k}:</b> {v}", meta_style))

    story.append(Spacer(1, 0.3 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1f4e79")))
    story.append(Spacer(1, 0.3 * cm))

    # ── Sections ──────────────────────────────────────────────────────────────
    HEADER_BG = colors.HexColor("#1f4e79")
    ROW_ALT = colors.HexColor("#eef2f7")

    for sec in sections:
        story.append(Paragraph(sec["title"], h2))
        if sec.get("notes"):
            story.append(Paragraph(sec["notes"], note_style))

        df: Optional[pd.DataFrame] = sec.get("df")
        if df is None or df.empty:
            story.append(Paragraph("No data.", meta_style))
            story.append(Spacer(1, 0.2 * cm))
            continue

        # Truncate long text cells to keep the table readable. astype(str)
        # doesn't reliably coerce every cell to `str` on all pandas
        # versions (e.g. NaN in some columns survives as float), so the
        # lambda coerces defensively rather than trusting astype alone.
        display_df = df.copy().astype(str).map(lambda x: str(x)[:120] if len(str(x)) > 120 else str(x))

        col_headers = list(display_df.columns)
        data_rows = display_df.values.tolist()
        table_data = [col_headers] + data_rows

        tbl = Table(table_data, repeatRows=1, hAlign="LEFT")
        tbl.setStyle(
            TableStyle(
                [
                    # Header row
                    ("BACKGROUND", (0, 0), (-1, 0), HEADER_BG),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, 0), 8),
                    # Data rows
                    ("FONTSIZE", (0, 1), (-1, -1), 7),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, ROW_ALT]),
                    # Grid
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#c0c0c0")),
                    ("LINEBELOW", (0, 0), (-1, 0), 0.5, colors.HexColor("#1f4e79")),
                    # Padding
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(tbl)
        story.append(Spacer(1, 0.4 * cm))

    doc.build(story)
    return buf.getvalue()


# ── DOCX ──────────────────────────────────────────────────────────────────────

def generate_docx(sections: List[ReportSection], meta: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    HEADER_BLUE = RGBColor(0x1F, 0x4E, 0x79)

    def _shade_cell(cell, hex_color: str) -> None:
        shd = cell._tc.get_or_add_tcPr()
        el = shd.makeelement(qn("w:shd"), {qn("w:fill"): hex_color})
        shd.append(el)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Pt(28)
    section.right_margin = Pt(28)

    title = doc.add_heading("Telecom Analyzer — Anomaly Detection Report", level=1)
    title.runs[0].font.color.rgb = HEADER_BLUE

    gen_p = doc.add_paragraph()
    gen_p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    for k, v in meta.items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(str(v))

    for sec in sections:
        doc.add_heading(sec["title"], level=2)
        if sec.get("notes"):
            note_p = doc.add_paragraph()
            note_p.add_run(sec["notes"]).italic = True

        df: Optional[pd.DataFrame] = sec.get("df")
        if df is None or df.empty:
            doc.add_paragraph("No data.")
            continue

        display_df = df.copy().astype(str).map(lambda x: str(x)[:200] if len(str(x)) > 200 else str(x))
        cols = list(display_df.columns)

        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(cols):
            hdr_cells[i].text = str(col)
            for p in hdr_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(8)
            _shade_cell(hdr_cells[i], "1F4E79")

        for _, row in display_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(cols):
                cells[i].text = str(row[col])
                for p in cells[i].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HTML ──────────────────────────────────────────────────────────────────────

def generate_html(
    sections: List[ReportSection],
    meta: Dict[str, Any],
    figures: Optional[list] = None,
    anomaly_cards: Optional[list] = None,
) -> bytes:
    """
    Generate a self-contained HTML report that mirrors the full dashboard.
    figures       : list of {"title": str, "fig": plotly Figure}
    anomaly_cards : list of {"severity", "title", "evidence",
                    "recommendation", "value", "unit", "warning", "critical"}
    """
    import plotly.io as pio

    _CSS = """
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Segoe UI',-apple-system,Arial,sans-serif;background:#f0f4f8;color:#1a1a2e}
.hdr{background:linear-gradient(135deg,#1f4e79 0%,#2d6ea8 100%);color:#fff;padding:2rem 2.5rem}
.hdr h1{font-size:1.75rem;font-weight:700;margin-bottom:.4rem}
.hdr .sub{opacity:.85;font-size:.88rem;margin-bottom:.9rem}
.hdr .chips{display:flex;flex-wrap:wrap;gap:.5rem}
.chip{background:rgba(255,255,255,.18);border-radius:6px;padding:.35rem .7rem;font-size:.8rem}
.wrap{max-width:1400px;margin:0 auto;padding:1.5rem}
.mrow{display:grid;grid-template-columns:repeat(auto-fill,minmax(150px,1fr));gap:1rem;margin:1.5rem 0}
.mc{background:#fff;border-radius:10px;padding:1.1rem;box-shadow:0 2px 8px rgba(0,0,0,.07);text-align:center;border-top:3px solid #1f4e79}  # noqa: E501
.mc .lbl{font-size:.7rem;color:#6c757d;text-transform:uppercase;letter-spacing:.06em;font-weight:600}
.mc .val{font-size:1.5rem;font-weight:700;color:#1f4e79;margin-top:.25rem}
.sec{background:#fff;border-radius:10px;padding:1.4rem;margin:1.1rem 0;box-shadow:0 2px 8px rgba(0,0,0,.07)}
.sec-title{font-size:1.05rem;font-weight:700;color:#1f4e79;padding-bottom:.55rem;margin-bottom:1rem;border-bottom:2px solid #1f4e79}  # noqa: E501
.notes{background:#e8f4fd;border-left:4px solid #1f4e79;padding:.45rem .75rem;border-radius:0 6px 6px 0;font-size:.83rem;color:#495057;margin-bottom:.9rem}  # noqa: E501
.tw{overflow-x:auto}
table{width:100%;border-collapse:collapse;font-size:.81rem}
thead tr{background:#1f4e79;color:#fff}
th{padding:8px 11px;text-align:left;font-weight:600;white-space:nowrap}
td{padding:6px 11px;border-bottom:1px solid #e9ecef;vertical-align:top;word-break:break-word;max-width:320px}
tbody tr:nth-child(even){background:#f0f4f8}
tbody tr:hover{background:#dde9f5}
.b{display:inline-block;padding:2px 8px;border-radius:12px;font-size:.74rem;font-weight:700}
.bc{background:#ffe0e0;color:#c0392b}
.bh{background:#ffecd2;color:#e67e22}
.bm{background:#fff9cc;color:#c9a000}
.bl{background:#d5f5e3;color:#27ae60}
.agrid{display:grid;grid-template-columns:repeat(auto-fill,minmax(440px,1fr));gap:1rem;margin-top:.8rem}
.acard{border-radius:8px;border:1px solid #dee2e6;overflow:hidden}
.acard-hdr{padding:.55rem 1rem;font-weight:700;font-size:.85rem;color:#fff}
.acard-hdr.critical{background:#c0392b}
.acard-hdr.high{background:#e67e22}
.acard-hdr.medium{background:#c9a000}
.acard-hdr.low{background:#27ae60}
.acard-body{padding:.8rem 1rem;display:grid;grid-template-columns:1fr 1fr;gap:.75rem;font-size:.8rem}
.ev{background:#e8f4fd;border-radius:6px;padding:.55rem}
.rc{background:#d5f5e3;border-radius:6px;padding:.55rem}
.albl{font-weight:700;font-size:.7rem;text-transform:uppercase;color:#6c757d;margin-bottom:.3rem}
.pbtn{position:fixed;bottom:1.8rem;right:1.8rem;background:#1f4e79;color:#fff;border:none;
      padding:.75rem 1.5rem;border-radius:8px;cursor:pointer;font-size:.9rem;
      box-shadow:0 4px 12px rgba(0,0,0,.25);z-index:9999;font-weight:600}
.pbtn:hover{background:#163d61}
.footer{font-size:.76rem;color:#999;text-align:right;margin-top:2rem;padding-top:1rem;border-top:1px solid #dee2e6}
@media print{
  .pbtn{display:none}
  body{background:#fff}
  .sec,.mc{box-shadow:none;border:1px solid #dee2e6}
  .hdr{print-color-adjust:exact;-webkit-print-color-adjust:exact}
  thead tr{print-color-adjust:exact;-webkit-print-color-adjust:exact}
  .acard-hdr{print-color-adjust:exact;-webkit-print-color-adjust:exact}
}
</style>
"""

    SEV_BADGE = {
        "Critical": '<span class="b bc">🔴 Critical</span>',
        "High": '<span class="b bh">🟠 High</span>',
        "Medium": '<span class="b bm">🟡 Medium</span>',
        "Low": '<span class="b bl">🟢 Low</span>',
    }

    p: List[str] = []

    # ── <head> ────────────────────────────────────────────────────────────────
    p.append(f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Telecom Analyzer Report — {meta.get('Source File', '')}</title>
{_CSS}""")
    if figures:
        p.append('<script src="https://cdn.plot.ly/plotly-latest.min.js"></script>\n')
    p.append("</head>\n<body>\n")

    # ── Header ────────────────────────────────────────────────────────────────
    chips = "".join(f'<span class="chip"><b>{k}:</b> {v}</span>' for k, v in meta.items())
    p.append(f"""<div class="hdr">
  <h1>📡 Telecom Analyzer — Anomaly Detection Report</h1>
  <div class="sub">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
  <div class="chips">{chips}</div>
</div>
<div class="wrap">
""")

    # ── Metric cards ──────────────────────────────────────────────────────────
    numeric = {k: v for k, v in meta.items()
               if isinstance(v, (int, float)) or str(v).replace(",", "").isdigit()}
    if numeric:
        p.append('<div class="mrow">')
        for k, v in numeric.items():
            p.append(f'<div class="mc"><div class="lbl">{k}</div><div class="val">{v}</div></div>')
        p.append('</div>\n')

    # ── Table sections ────────────────────────────────────────────────────────
    for sec in sections:
        df: Optional[pd.DataFrame] = sec.get("df")
        p.append(f'<div class="sec"><div class="sec-title">{sec["title"]}</div>')
        if sec.get("notes"):
            p.append(f'<div class="notes">{sec["notes"]}</div>')
        if df is not None and not df.empty:
            p.append('<div class="tw"><table><thead><tr>')
            for col in df.columns:
                p.append(f'<th>{col}</th>')
            p.append('</tr></thead><tbody>')
            for _, row in df.iterrows():
                p.append('<tr>')
                for col in df.columns:
                    val = str(row[col])
                    if col == "Severity":
                        p.append(f'<td>{SEV_BADGE.get(val, val)}</td>')
                    else:
                        p.append(f'<td>{val}</td>')
                p.append('</tr>')
            p.append('</tbody></table></div>')
        else:
            p.append('<p style="color:#6c757d;padding:.5rem 0">No data.</p>')
        p.append('</div>\n')

    # ── Plotly charts ─────────────────────────────────────────────────────────
    if figures:
        for item in figures:
            chart_html = pio.to_html(item["fig"], include_plotlyjs=False, full_html=False)
            p.append(f'<div class="sec"><div class="sec-title">{item["title"]}</div>{chart_html}</div>\n')

    # ── Anomaly detail cards ──────────────────────────────────────────────────
    if anomaly_cards:
        p.append('<div class="sec"><div class="sec-title">🔎 Anomaly Detail Cards</div><div class="agrid">')
        for card in anomaly_cards:
            sev_lower = card["severity"].lower()
            badge = SEV_BADGE.get(card["severity"], card["severity"])
            ev = str(card.get("evidence", "—")).replace("<", "&lt;").replace(">", "&gt;")
            rec = str(card.get("recommendation", "—")).replace("<", "&lt;").replace(">", "&gt;")
            p.append(f"""<div class="acard">
  <div class="acard-hdr {sev_lower}">{badge} &nbsp; {card["title"]}</div>
  <div class="acard-body">
    <div class="ev"><div class="albl">Evidence</div>{ev}
      <br><small><b>Value:</b> {card.get("value", "—")} {card.get("unit", "")}
      &nbsp;|&nbsp; Warn: {card.get("warning", "—")}
      &nbsp;|&nbsp; Crit: {card.get("critical", "—")}</small></div>
    <div class="rc"><div class="albl">Recommendation</div>{rec}</div>
  </div>
</div>""")
        p.append('</div></div>\n')

    # ── Footer ────────────────────────────────────────────────────────────────
    p.append(f'<div class="footer">Telecom Analyzer · {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</div>')
    p.append("""
<button class="pbtn" onclick="window.print()">🖨️ Print / Save as PDF</button>
</div>
</body>
</html>""")

    return "".join(p).encode("utf-8")
